from __future__ import annotations

"""
Document parser — extracts structured text from PDF, DOCX, and Excel files.
Returns a list of (page_number, text) tuples to preserve page provenance.
"""

from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook


def parse_pdf(filepath: str) -> list[tuple[int, str]]:
    """Extract text per page from a PDF file using pdfplumber."""
    pages: list[tuple[int, str]] = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append((i, text))
    return pages


def parse_docx(filepath: str) -> list[tuple[int, str]]:
    """
    Extract text from a DOCX file.
    DOCX doesn't have native page numbers, so we synthesize page breaks
    based on paragraph count (~40 paragraphs ≈ 1 page).
    """
    doc = DocxDocument(filepath)
    full_paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_paragraphs.append(text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                full_paragraphs.append(" | ".join(cells))

    # Synthesize pages (~40 paragraphs per page)
    PAGE_SIZE = 40
    pages: list[tuple[int, str]] = []
    for i in range(0, len(full_paragraphs), PAGE_SIZE):
        page_num = (i // PAGE_SIZE) + 1
        page_text = "\n".join(full_paragraphs[i : i + PAGE_SIZE])
        if page_text.strip():
            pages.append((page_num, page_text))

    return pages


def parse_xlsx(filepath: str) -> list[tuple[int, str]]:
    """
    Extract text from an Excel (.xlsx) file.
    Each worksheet becomes a "page" with sheet_index as the page number.
    Rows are joined with pipe separators (consistent with DOCX table extraction).
    """
    wb = load_workbook(filepath, read_only=True, data_only=True)
    pages: list[tuple[int, str]] = []

    for sheet_num, sheet_name in enumerate(wb.sheetnames, start=1):
        ws = wb[sheet_name]
        rows: list[str] = []

        rows.append(f"Sheet: {sheet_name}")

        for row in ws.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None]
            if cells:
                rows.append(" | ".join(cells))

        page_text = "\n".join(rows)
        if page_text.strip() and len(rows) > 1:
            pages.append((sheet_num, page_text))

    wb.close()
    return pages


def parse_document(filepath: str) -> list[tuple[int, str]]:
    """Auto-detect format and parse."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return parse_docx(filepath)
    elif ext == ".xlsx":
        return parse_xlsx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
