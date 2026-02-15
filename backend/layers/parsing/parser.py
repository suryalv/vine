from __future__ import annotations

"""
PARSING LAYER
=============
Extracts structured text from PDF, DOCX, and Excel files.
Returns list of (page_number, text) tuples to preserve page provenance.

Team Owner: Document Ingestion Team
"""

from datetime import date, datetime
from pathlib import Path
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook


def parse_pdf(filepath: str) -> list[tuple[int, str]]:
    """Extract text per page from a PDF file using pdfplumber."""
    pages: list[tuple[int, str]] = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append((i, text))
    return pages


def parse_docx(filepath: str) -> list[tuple[int, str]]:
    """
    Extract text from a DOCX file.
    DOCX doesn't have native page numbers, so we synthesize page breaks
    based on paragraph count (~40 paragraphs = 1 page).
    """
    doc = DocxDocument(filepath)
    full_paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_paragraphs.append(text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                full_paragraphs.append(" | ".join(cells))

    # Synthesize pages (~40 paragraphs per page)
    PAGE_SIZE = 40
    pages: list[tuple[int, str]] = []
    for i in range(0, len(full_paragraphs), PAGE_SIZE):
        page_num = (i // PAGE_SIZE) + 1
        page_text = "\n".join(full_paragraphs[i : i + PAGE_SIZE])
        if page_text.strip():
            pages.append((page_num, page_text))

    return pages


def _format_cell(value) -> str:
    """Format a cell value for text output. Returns '' for None or formula strings."""
    if value is None:
        return ""
    if isinstance(value, datetime):
        if value.hour == 0 and value.minute == 0 and value.second == 0:
            return value.strftime("%Y-%m-%d")
        return value.strftime("%Y-%m-%d %H:%M")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    text = str(value).strip()
    # Strip raw formula strings — they're noise for RAG/embeddings
    if text.startswith("="):
        return ""
    text = text.replace("\n", " ").replace("\r", " ")
    return text


# Large sheets are split into chunks to keep text sizes manageable for embeddings
_XLSX_ROWS_PER_PAGE = 100


def parse_xlsx(filepath: str) -> list[tuple[int, str]]:
    """
    Extract text from an Excel (.xlsx) file.
    Each worksheet becomes one or more "pages". Large sheets (>100 rows)
    are split into chunks with the header row repeated in each chunk.
    Rows are pipe-separated (consistent with DOCX tables).
    Hidden sheets are skipped.
    """
    wb = load_workbook(filepath, read_only=True, data_only=True)
    pages: list[tuple[int, str]] = []
    page_counter = 1

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # Skip hidden sheets
        if ws.sheet_state != "visible":
            continue

        all_rows: list[str] = []
        header_line = f"Sheet: {sheet_name}"
        header_row: str | None = None
        seen_values: set[str] = set()

        for row in ws.iter_rows(values_only=True):
            cells = [_format_cell(cell) for cell in row]
            cells = [c for c in cells if c]
            if not cells:
                continue
            line = " | ".join(cells)
            # Deduplicate rows caused by merged cells
            if line in seen_values:
                continue
            seen_values.add(line)
            # Capture first data row as column header for chunk repetition
            if header_row is None:
                header_row = line
            all_rows.append(line)

        if not all_rows:
            continue

        # Split large sheets into chunks, repeat header row in each chunk
        for chunk_start in range(0, len(all_rows), _XLSX_ROWS_PER_PAGE):
            chunk = all_rows[chunk_start : chunk_start + _XLSX_ROWS_PER_PAGE]
            chunk_label = header_line
            if len(all_rows) > _XLSX_ROWS_PER_PAGE:
                chunk_label += f" (rows {chunk_start + 1}-{chunk_start + len(chunk)})"
            lines = [chunk_label]
            # Repeat header row in continuation chunks for context
            if chunk_start > 0 and header_row:
                lines.append(header_row)
            lines.extend(chunk)
            pages.append((page_counter, "\n".join(lines)))
            page_counter += 1

    wb.close()
    return pages


def parse_document(filepath: str) -> list[tuple[int, str]]:
    """Auto-detect format and parse."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return parse_docx(filepath)
    elif ext in (".xlsx", ".xls"):
        return parse_xlsx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
