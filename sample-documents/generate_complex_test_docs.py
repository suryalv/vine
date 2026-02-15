#!/usr/bin/env python3
"""
Generate highly complex underwriting test documents.
These documents are intentionally dense and interrelated to stress-test
the RAG pipeline's parsing, chunking, embedding, retrieval, and
hallucination detection capabilities.
"""

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─── reportlab styles ─────────────────────────────────────────
styles = getSampleStyleSheet()
styles.add(ParagraphStyle(name='DocTitle', fontSize=18, leading=22, spaceAfter=6,
           fontName='Helvetica-Bold', textColor=colors.HexColor('#00205B')))
styles.add(ParagraphStyle(name='DocSubtitle', fontSize=11, leading=14, spaceAfter=12,
           textColor=colors.HexColor('#475569')))
styles.add(ParagraphStyle(name='SectionHead', fontSize=13, leading=16, spaceBefore=18,
           spaceAfter=8, fontName='Helvetica-Bold', textColor=colors.HexColor('#0f172a')))
styles.add(ParagraphStyle(name='SubSection', fontSize=11, leading=14, spaceBefore=12,
           spaceAfter=6, fontName='Helvetica-Bold', textColor=colors.HexColor('#1e293b')))
styles.add(ParagraphStyle(name='BodyText2', fontSize=9.5, leading=13, spaceAfter=6,
           textColor=colors.HexColor('#334155')))
styles.add(ParagraphStyle(name='SmallGray', fontSize=8, leading=10,
           textColor=colors.HexColor('#94a3b8')))
styles.add(ParagraphStyle(name='CenterBold', fontSize=10, leading=13,
           fontName='Helvetica-Bold', alignment=TA_CENTER))
styles.add(ParagraphStyle(name='Footer', fontSize=7.5, leading=10,
           textColor=colors.HexColor('#94a3b8'), alignment=TA_CENTER))

AIG_BLUE = colors.HexColor('#00205B')
LIGHT_BLUE = colors.HexColor('#e8edf5')
BORDER_GRAY = colors.HexColor('#cbd5e1')
ROW_ALT = colors.HexColor('#f8fafc')
RED_FLAG = colors.HexColor('#dc2626')


def header_table(title_lines):
    data = [
        ['AIG', title_lines[0]],
        ['', title_lines[1] if len(title_lines) > 1 else ''],
    ]
    t = Table(data, colWidths=[1.2 * inch, 5.3 * inch])
    t.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (0, 0), 20),
        ('TEXTCOLOR', (0, 0), (0, 0), AIG_BLUE),
        ('FONTNAME', (1, 0), (1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (1, 0), (1, 0), 14),
        ('TEXTCOLOR', (1, 0), (1, 0), AIG_BLUE),
        ('FONTSIZE', (1, 1), (1, 1), 9),
        ('TEXTCOLOR', (1, 1), (1, 1), colors.HexColor('#64748b')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
    ]))
    return t


def styled_table(headers, rows, col_widths=None):
    data = [headers] + rows
    if not col_widths:
        col_widths = [6.5 * inch / len(headers)] * len(headers)
    t = Table(data, colWidths=col_widths, repeatRows=1)
    style = [
        ('BACKGROUND', (0, 0), (-1, 0), AIG_BLUE),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 8.5),
        ('FONTSIZE', (0, 1), (-1, -1), 8.5),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, BORDER_GRAY),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
    ]
    for i in range(1, len(data)):
        if i % 2 == 0:
            style.append(('BACKGROUND', (0, i), (-1, i), ROW_ALT))
    t.setStyle(TableStyle(style))
    return t


def kv_table(pairs, col_widths=None):
    if not col_widths:
        col_widths = [2.2 * inch, 4.3 * inch]
    t = Table(pairs, colWidths=col_widths)
    style = [
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#1e293b')),
        ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor('#334155')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LINEBELOW', (0, 0), (-1, -2), 0.5, colors.HexColor('#e2e8f0')),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
    ]
    t.setStyle(TableStyle(style))
    return t


def add_footer(canvas, doc):
    canvas.saveState()
    canvas.setFont('Helvetica', 7)
    canvas.setFillColor(colors.HexColor('#94a3b8'))
    canvas.drawCentredString(4.25 * inch, 0.4 * inch,
                             f"AIG  |  Confidential  |  Page {doc.page}")
    canvas.restoreState()


# ─── docx helpers ──────────────────────────────────────────────
AIG_BLUE_RGB = RGBColor(0x00, 0x20, 0x5B)
BODY_TEXT_RGB = RGBColor(0x33, 0x41, 0x55)
MUTED_RGB = RGBColor(0x64, 0x74, 0x8B)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def style_header_row(row, bg_color='00205B'):
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE_RGB
                run.font.bold = True
                run.font.size = Pt(9)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = AIG_BLUE_RGB
    return h


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT_RGB
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


# ============================================================
# 1. WORKERS' COMPENSATION — MULTI-STATE RETROSPECTIVE POLICY
#    (PDF — 6 pages, very dense)
# ============================================================
def generate_workers_comp_retro():
    path = os.path.join(OUT_DIR, 'Workers_Comp_Retro_Rating_Atlas_Industries.pdf')
    doc = SimpleDocTemplate(path, pagesize=letter,
                            topMargin=0.6 * inch, bottomMargin=0.7 * inch,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    story = []

    story.append(header_table([
        "Workers' Compensation & Employers' Liability",
        'Retrospective Rating Plan  |  Policy No: AIG-WC-2026-RR-00412'
    ]))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=2, color=AIG_BLUE))
    story.append(Spacer(1, 14))

    # ── Declarations ──
    story.append(Paragraph('DECLARATIONS', styles['SectionHead']))
    story.append(kv_table([
        ['Named Insured:', 'Atlas Industries, Inc. and all subsidiaries'],
        ['Subsidiaries:', 'Atlas Steel Fabrication, LLC\nAtlas Heavy Equipment Rental, LLC\nAtlas Environmental Services, Inc.\nAtlas Marine Construction, LLC'],
        ['Mailing Address:', '1200 Commerce Tower, 15th Floor\nNew Orleans, LA 70130'],
        ['FEIN:', '72-8841520 (Parent)'],
        ['Policy Period:', '01/01/2026 to 01/01/2027 — 12:01 AM Standard Time'],
        ['Rating Plan:', 'Retrospective Rating — Large Deductible ($250,000 per occurrence)'],
        ['Monopolistic States:', 'Separate policies required: OH, WA, WY, ND (excluded from this policy)'],
        ['Premium Basis:', 'Estimated Annual Premium: $2,847,000 (subject to retrospective adjustment)'],
        ['Deposit Premium:', '$1,423,500 (50% due at inception)'],
        ['Producer:', 'Willis Towers Watson — Energy & Construction Practice\nAttn: Robert Callahan, Managing Director'],
    ]))
    story.append(Spacer(1, 10))

    # ── Part One: WC ──
    story.append(Paragraph('PART ONE — WORKERS\' COMPENSATION INSURANCE', styles['SectionHead']))
    story.append(Paragraph('Coverage applies in the following states for the class codes and estimated payrolls indicated. '
                           'Atlas Industries operates in 14 states with a workforce of 2,847 employees across 28 work sites.',
                           styles['BodyText2']))
    story.append(Spacer(1, 6))

    story.append(styled_table(
        ['State', 'Class\nCode', 'Description', 'Est. Annual\nPayroll', 'Rate\nper $100', 'Est.\nPremium', 'Employees'],
        [
            ['LA', '5057', 'Iron/Steel Erection — structural', '$8,420,000', '$14.82', '$1,247,844', '186'],
            ['LA', '3620', 'Steel Fabrication Shop', '$4,280,000', '$7.24', '$309,872', '124'],
            ['LA', '8810', 'Clerical Office Employees', '$2,140,000', '$0.18', '$3,852', '42'],
            ['TX', '5057', 'Iron/Steel Erection — structural', '$6,840,000', '$12.47', '$852,948', '148'],
            ['TX', '5022', 'Masonry Work — NOC', '$2,120,000', '$8.91', '$188,892', '52'],
            ['TX', '7219', 'Trucking — heavy equipment hauling', '$1,680,000', '$11.34', '$190,512', '38'],
            ['TX', '8810', 'Clerical Office', '$1,240,000', '$0.14', '$1,736', '28'],
            ['MS', '5057', 'Iron/Steel Erection', '$3,240,000', '$16.28', '$527,472', '74'],
            ['MS', '6217', 'Excavation & Grading', '$2,180,000', '$9.45', '$206,010', '56'],
            ['AL', '5057', 'Iron/Steel Erection', '$2,860,000', '$13.71', '$392,106', '62'],
            ['AL', '3620', 'Steel Fabrication Shop', '$1,940,000', '$6.89', '$133,666', '48'],
            ['FL', '5057', 'Iron/Steel Erection', '$4,120,000', '$18.92', '$779,504', '92'],
            ['FL', '5403', 'Marine Construction — pile driving', '$3,680,000', '$22.14', '$814,752', '84'],
            ['FL', '6251', 'Tunneling — soft ground', '$1,480,000', '$24.87', '$368,076', '32'],
            ['GA', '5057', 'Iron/Steel Erection', '$1,920,000', '$11.84', '$227,328', '44'],
            ['GA', '8227', 'Equipment Rental — w/operators', '$2,480,000', '$8.42', '$208,816', '58'],
            ['SC', '5057', 'Iron/Steel Erection', '$1,640,000', '$12.56', '$205,984', '38'],
            ['TN', '5057', 'Iron/Steel Erection', '$1,280,000', '$10.94', '$140,032', '28'],
            ['NC', '3724', 'Structural Steel Fabrication', '$2,840,000', '$5.67', '$161,028', '68'],
            ['VA', '5057', 'Iron/Steel Erection', '$980,000', '$13.21', '$129,458', '22'],
            ['AR', '5057', 'Iron/Steel Erection', '$720,000', '$15.43', '$111,096', '16'],
            ['KY', '5057', 'Iron/Steel Erection', '$480,000', '$11.78', '$56,544', '12'],
        ],
        col_widths=[0.4 * inch, 0.45 * inch, 1.6 * inch, 0.9 * inch, 0.65 * inch, 0.85 * inch, 0.65 * inch]
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph('<b>Total Estimated Payroll:</b> $58,540,000  |  <b>Total Estimated Manual Premium:</b> $7,708,528  |  <b>Total Employees:</b> 1,252 (shown above; 1,595 in excluded monopolistic states)',
                           styles['BodyText2']))

    story.append(PageBreak())

    # ── Part Two: EL ──
    story.append(Paragraph("PART TWO — EMPLOYERS' LIABILITY", styles['SectionHead']))
    story.append(kv_table([
        ['Coverage A:', 'Workers\' Compensation — Statutory'],
        ['Coverage B — Bodily Injury by Accident:', '$1,000,000 each accident'],
        ['Coverage B — Bodily Injury by Disease:', '$1,000,000 policy limit'],
        ['Coverage B — Bodily Injury by Disease:', '$1,000,000 each employee'],
        ['Stop Gap Coverage (Monopolistic States):', '$1,000,000 — OH, WA, WY, ND'],
        ['Maritime / USL&H Coverage:', 'Included — FL, LA, TX, MS, AL\nPer 33 USC §901 et seq.'],
        ['Voluntary Compensation:', 'Included for Executive Officers (Sub: Atlas Industries only)'],
        ['Foreign Voluntary WC:', 'Coverage Territory: Mexico, Canada, Caribbean\n$1,000,000 each accident'],
    ]))
    story.append(Spacer(1, 10))

    # ── Experience Mod ──
    story.append(Paragraph('EXPERIENCE MODIFICATION', styles['SectionHead']))
    story.append(Paragraph('The following Experience Modification Rates (EMR) are effective for the current policy period. '
                           'EMRs are calculated by NCCI (or applicable state rating bureau) based on 3 years of loss history '
                           'excluding the most recent year.',
                           styles['BodyText2']))
    story.append(Spacer(1, 6))

    story.append(styled_table(
        ['State', 'Bureau', 'Effective\nDate', 'EMR', 'Expected\nLosses', 'Actual\nLosses', 'Ballast\nValue', 'Primary\nWeight'],
        [
            ['LA', 'NCCI', '01/01/2026', '1.18', '$412,000', '$486,300', '$148,200', '0.42'],
            ['TX', 'NCCI', '01/01/2026', '0.97', '$284,000', '$275,480', '$102,400', '0.38'],
            ['MS', 'NCCI', '01/01/2026', '1.24', '$186,000', '$230,640', '$67,100', '0.35'],
            ['AL', 'NCCI', '01/01/2026', '1.08', '$124,000', '$133,920', '$44,600', '0.33'],
            ['FL', 'NCCI', '01/01/2026', '1.32', '$398,000', '$525,360', '$143,200', '0.44'],
            ['GA', 'NCCI', '01/01/2026', '0.92', '$98,000', '$90,160', '$35,200', '0.30'],
            ['SC', 'NCCI', '01/01/2026', '1.05', '$62,000', '$65,100', '$22,300', '0.28'],
            ['TN', 'NCCI', '01/01/2026', '0.89', '$44,000', '$39,160', '$15,800', '0.26'],
            ['NC', 'NCRIB', '01/01/2026', '1.02', '$48,000', '$48,960', '$17,200', '0.27'],
            ['VA', 'NCCI', '01/01/2026', '0.94', '$32,000', '$30,080', '$11,500', '0.24'],
        ],
        col_widths=[0.4 * inch, 0.5 * inch, 0.65 * inch, 0.45 * inch, 0.7 * inch, 0.7 * inch, 0.65 * inch, 0.55 * inch]
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph('<b>Weighted Average EMR:</b> 1.14  |  <b>Industry Average:</b> 1.00  |  '
                           '<b>EMR Trend:</b> Increasing (was 1.08 in 2024, 1.11 in 2025)',
                           styles['BodyText2']))
    story.append(Paragraph('<i>NOTE: Florida EMR of 1.32 is driven by 2 large indemnity claims (CLM-2023-FL-001, CLM-2024-FL-002) '
                           'involving marine pile driving operations. See loss detail below.</i>',
                           styles['BodyText2']))

    story.append(PageBreak())

    # ── Retrospective Rating ──
    story.append(Paragraph('RETROSPECTIVE RATING PLAN PARAMETERS', styles['SectionHead']))
    story.append(Paragraph('This policy is subject to retrospective premium adjustment based on actual incurred losses '
                           'during the policy period. The following parameters govern the retrospective calculation:',
                           styles['BodyText2']))
    story.append(Spacer(1, 6))
    story.append(kv_table([
        ['Plan Type:', 'Incurred Loss Retro — Annual Adjustments for 5 years post-expiration'],
        ['Basic Premium Factor:', '0.286 (includes tax multiplier, expense constant, carrier risk charge)'],
        ['Loss Conversion Factor (LCF):', '1.12 (includes ALAE loading)'],
        ['Tax Multiplier:', '1.042 (weighted average across all states)'],
        ['Per-Occurrence Limitation:', '$250,000 (losses capped at this amount per occurrence for retro calc)'],
        ['Minimum Retrospective Premium:', '$1,994,900 (70% of Standard Premium)'],
        ['Maximum Retrospective Premium:', '$4,274,700 (150% of Standard Premium)'],
        ['Standard Premium:', '$2,849,800 (Manual Premium × EMR × Schedule Credit/Debit)'],
        ['Schedule Rating Credit:', '-8% (approved by NCCI effective 01/01/2026)'],
        ['Retro Development Period:', '1st adjustment: 6 months post-expiration, then annually × 5 years'],
        ['Collateral Requirement:', 'Irrevocable Letter of Credit: $1,200,000\n(subject to annual review based on loss development)'],
    ]))
    story.append(Spacer(1, 10))

    story.append(Paragraph('RETRO FORMULA', styles['SubSection']))
    story.append(Paragraph('Retrospective Premium = (Basic Premium + Converted Losses) × Tax Multiplier',
                           styles['CenterBold']))
    story.append(Spacer(1, 4))
    story.append(Paragraph('Where: Converted Losses = (Capped Incurred Losses) × Loss Conversion Factor (1.12)',
                           styles['BodyText2']))
    story.append(Paragraph('Subject to: Minimum Premium ≤ Retro Premium ≤ Maximum Premium',
                           styles['BodyText2']))
    story.append(Spacer(1, 10))

    # ── Premium Computation ──
    story.append(Paragraph('PREMIUM COMPUTATION SUMMARY', styles['SectionHead']))
    story.append(styled_table(
        ['Component', 'Amount', 'Notes'],
        [
            ['Total Manual Premium', '$7,708,528', 'Sum of all state/class premiums'],
            ['Experience Modification (1.14)', '×1.14', 'Weighted average EMR'],
            ['Modified Premium', '$8,787,722', 'Manual × EMR'],
            ['Schedule Rating Credit', '-8%', 'NCCI approved'],
            ['Scheduled Premium', '$8,084,704', 'Modified × 0.92'],
            ['Premium Discount', '-12.4%', 'Per NCCI Table (>$5M tier)'],
            ['Standard Premium', '$7,082,201', 'Base for retro calculation'],
            ['Expense Constant', '$960', '$160 × 6 entities'],
            ['TRIA Surcharge', '$28,400', '0.4% of Standard Premium'],
            ['Catastrophe Provision', '$42,600', '0.6% of Standard Premium'],
            ['Estimated Annual Premium', '$2,847,000', 'Deposit (retro adj. follows)'],
        ],
        col_widths=[2.0 * inch, 1.2 * inch, 3.3 * inch]
    ))

    story.append(PageBreak())

    # ── Loss History ──
    story.append(Paragraph('5-YEAR LOSS HISTORY — ALL SUBSIDIARIES', styles['SectionHead']))
    story.append(Paragraph('Source: AIG Claims Division, valued as of 12/31/2025. Includes all ALAE. '
                           'Experience period used for EMR: 01/01/2022 — 12/31/2024.',
                           styles['SmallGray']))
    story.append(Spacer(1, 6))

    story.append(styled_table(
        ['Year', 'Entity', 'Claims\nCount', 'Open', 'Medical\nPaid', 'Indemnity\nPaid', 'Reserves', 'Total\nIncurred'],
        [
            ['2025', 'Atlas Steel Fabrication', '28', '8', '$142,800', '$284,600', '$186,000', '$613,400'],
            ['2025', 'Atlas Heavy Equipment', '14', '4', '$68,200', '$124,000', '$92,000', '$284,200'],
            ['2025', 'Atlas Environmental', '8', '2', '$34,600', '$48,200', '$28,000', '$110,800'],
            ['2025', 'Atlas Marine Construction', '22', '7', '$186,400', '$342,800', '$248,000', '$777,200'],
            ['2024', 'Atlas Steel Fabrication', '32', '2', '$168,400', '$312,000', '$24,000', '$504,400'],
            ['2024', 'Atlas Heavy Equipment', '18', '0', '$82,400', '$148,600', '$0', '$231,000'],
            ['2024', 'Atlas Environmental', '6', '0', '$22,800', '$38,400', '$0', '$61,200'],
            ['2024', 'Atlas Marine Construction', '26', '1', '$218,600', '$486,200', '$42,000', '$746,800'],
            ['2023', 'Atlas Steel Fabrication', '24', '0', '$124,200', '$218,400', '$0', '$342,600'],
            ['2023', 'Atlas Marine Construction', '18', '0', '$142,000', '$384,600', '$0', '$526,600'],
            ['2023', 'All Other Entities', '22', '0', '$86,400', '$128,200', '$0', '$214,600'],
            ['2022', 'All Entities Combined', '64', '0', '$284,000', '$412,800', '$0', '$696,800'],
            ['2021', 'All Entities Combined', '52', '0', '$218,600', '$328,400', '$0', '$547,000'],
        ],
        col_widths=[0.45 * inch, 1.5 * inch, 0.5 * inch, 0.4 * inch, 0.7 * inch, 0.75 * inch, 0.65 * inch, 0.7 * inch]
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph('<b>5-Year Total Incurred:</b> $5,656,600  |  <b>Avg Annual Incurred:</b> $1,131,320  |  '
                           '<b>Total Claims:</b> 334  |  <b>Open Claims:</b> 24',
                           styles['BodyText2']))
    story.append(Spacer(1, 10))

    # ── Large Loss Detail ──
    story.append(Paragraph('LARGE LOSS DETAIL (>$100,000 INCURRED)', styles['SectionHead']))
    large_losses = [
        ['CLM-2024-FL-002', '08/14/2024', 'Atlas Marine\nConstruction', 'FL — 5403',
         'Pile driving barge capsized during Hurricane Debby approach. 2 workers fell into water. '
         'Worker A: bilateral leg fractures, 8 months disability. Worker B: traumatic brain injury, '
         'permanent partial disability. OSHA Citation #2024-FL-08841 issued for inadequate weather '
         'monitoring protocol.', '$0', '$386,000', '$386,000', 'OPEN'],
        ['CLM-2023-FL-001', '03/22/2023', 'Atlas Marine\nConstruction', 'FL — 5403',
         'Steel beam dropped from crane during bridge reinforcement. Worker struck on shoulder — '
         'rotator cuff tear, cervical disc herniation. Surgery × 2. TTD 14 months. MMI reached '
         '09/2024 with 22% whole person impairment rating.', '$218,400', '$164,200', '$382,600', 'CLOSED'],
        ['CLM-2025-LA-001', '02/18/2025', 'Atlas Steel\nFabrication', 'LA — 5057',
         'Fall from height (42 ft) during structural steel erection at refinery turnaround project. '
         'Worker suffered multiple rib fractures, pneumothorax, and L1 compression fracture. Currently '
         'in rehabilitation. SIF (Serious Injury/Fatality) potential flagged.', '$84,200', '$124,000',
         '$208,200', 'OPEN'],
        ['CLM-2024-TX-001', '06/02/2024', 'Atlas Heavy\nEquipment', 'TX — 7219',
         'Equipment hauling truck overturned on I-10 near Beaumont. Driver ejected — multiple fractures, '
         'internal injuries. Subrogation potential against road construction contractor (wet road surface). '
         'Subro recovery of $42,000 credited.', '$148,600', '$0', '$148,600', 'CLOSED'],
        ['CLM-2025-MS-001', '07/11/2025', 'Atlas Environmental\nServices', 'MS — 6217',
         'Trench collapse during environmental remediation excavation. Worker buried to waist — bilateral '
         'femur fractures. OSHA Willful Citation issued (inadequate shoring, no competent person on site). '
         'Penalty: $156,000. Legal defense costs mounting.', '$34,600', '$118,000', '$152,600', 'OPEN'],
    ]

    for ll in large_losses:
        story.append(Paragraph(f'Claim: {ll[0]}', styles['SubSection']))
        story.append(kv_table([
            ['Date of Loss:', ll[1]],
            ['Entity:', ll[2]],
            ['State/Class:', ll[3]],
            ['Status:', ll[8]],
        ], col_widths=[1.2 * inch, 5.3 * inch]))
        story.append(Paragraph(f'<b>Description:</b> {ll[4]}', styles['BodyText2']))
        fin_data = [
            ['Medical + Indemnity Paid', 'Outstanding Reserves', 'Total Incurred'],
            [ll[5], ll[6], ll[7]],
        ]
        ft = Table(fin_data, colWidths=[2.1 * inch, 2.2 * inch, 2.2 * inch])
        ft.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BACKGROUND', (0, 0), (-1, 0), LIGHT_BLUE),
            ('GRID', (0, 0), (-1, -1), 0.5, BORDER_GRAY),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(ft)
        story.append(Spacer(1, 8))

    story.append(PageBreak())

    # ── Safety & Risk Control ──
    story.append(Paragraph('RISK CONTROL & SAFETY REQUIREMENTS', styles['SectionHead']))
    reqs = [
        '<b>1. OSHA Compliance:</b> Atlas Industries shall maintain an active safety program with a designated '
        'Director of Safety (currently: Michael Torres, CSP). All OSHA citations must be reported to AIG within '
        '10 business days of receipt. The current DART rate is 4.8 (industry average: 3.2) and TRIR is 6.4 '
        '(industry average: 4.1). AIG requires a 15% reduction in TRIR by renewal or a 5% surcharge will apply.',

        '<b>2. Marine Operations — Enhanced Requirements:</b> Given the adverse loss experience in Class 5403 '
        '(Marine Construction), the following additional requirements apply: (a) All marine operations must have '
        'a site-specific safety plan approved by a Certified Marine Chemist; (b) Weather monitoring protocol must '
        'include 48-hour advance storm tracking with mandatory cessation of overwater operations when sustained '
        'winds exceed 25 mph; (c) Quarterly safety audits by a third-party maritime safety consultant.',

        '<b>3. Fall Protection Program:</b> All work at heights greater than 6 feet (general industry) or 15 feet '
        '(steel erection per OSHA 29 CFR 1926.760) requires documented Job Hazard Analysis (JHA), properly inspected '
        'personal fall arrest systems, and rescue planning. Near-miss reports for fall events must be submitted to AIG '
        'Risk Control quarterly.',

        '<b>4. Trench & Excavation Safety:</b> Following the Willful Citation (CLM-2025-MS-001), Atlas Environmental '
        'Services must: (a) retain a full-time competent person for all trenching operations per 29 CFR 1926 Subpart P; '
        '(b) implement a trench safety audit checklist; (c) provide AIG with monthly compliance reports until further notice.',

        '<b>5. Return-to-Work Program:</b> Atlas shall maintain a formal Return-to-Work (RTW) program offering modified '
        'duty positions within 48 hours of injury. The current RTW participation rate is 62%. AIG targets a minimum 80% '
        'participation rate. A dedicated RTW coordinator must be assigned by 03/01/2026.',

        '<b>6. Fleet Safety:</b> All drivers in Class 7219 (Heavy Equipment Hauling) must complete annual defensive driving '
        'training, pre-trip inspection certification, and random drug/alcohol testing per DOT regulations. MVR reviews '
        'required semi-annually. Zero tolerance for DUI/DWI violations.',
    ]
    for r in reqs:
        story.append(Paragraph(r, styles['BodyText2']))
        story.append(Spacer(1, 4))

    story.append(Spacer(1, 14))
    story.append(HRFlowable(width="100%", thickness=1, color=BORDER_GRAY))
    story.append(Spacer(1, 6))
    story.append(Paragraph('This policy is issued subject to the retrospective rating plan endorsement (WC 00 05 05) '
                           'and all endorsements listed in the policy jacket. Premium adjustments will be computed in '
                           'accordance with the NCCI retrospective rating manual and applicable state rating bureau rules.',
                           styles['SmallGray']))

    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    print(f"  Created: {path}")


# ============================================================
# 2. EXCESS LIABILITY (UMBRELLA) — MULTI-LAYER TOWER
#    (PDF — 5 pages)
# ============================================================
def generate_excess_umbrella():
    path = os.path.join(OUT_DIR, 'Excess_Umbrella_MultiLayer_TechCorp_Global.pdf')
    doc = SimpleDocTemplate(path, pagesize=letter,
                            topMargin=0.6 * inch, bottomMargin=0.7 * inch,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    story = []

    story.append(header_table([
        'Commercial Excess Liability (Umbrella)',
        'Multi-Layer Tower  |  Policy No: AIG-XS-2026-GL-08842'
    ]))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", thickness=2, color=AIG_BLUE))
    story.append(Spacer(1, 14))

    # ── Declarations ──
    story.append(Paragraph('DECLARATIONS', styles['SectionHead']))
    story.append(kv_table([
        ['Named Insured:', 'TechCorp Global Holdings, Inc. and all subsidiaries\n'
                          'as defined in the Scheduled Underlying Policies'],
        ['DBA / Subsidiaries:', 'TechCorp Software Solutions, LLC\n'
                               'TechCorp Cloud Infrastructure, Inc.\n'
                               'TechCorp Autonomous Systems, LLC\n'
                               'TechCorp Financial Services, Ltd. (UK)\n'
                               'TechCorp Asia Pacific Pte. Ltd. (Singapore)'],
        ['Mailing Address:', '2000 Innovation Drive, Suite 4000\nSan Jose, CA 95134'],
        ['FEIN:', '94-7228341 (Parent)'],
        ['Policy Period:', '06/01/2026 to 06/01/2027 — 12:01 AM Pacific Time'],
        ['Policy Form:', 'AIG Excess Follow-Form w/ Drop-Down (AIG-XS-2026)'],
        ['Retention:', 'Per underlying scheduled limits (see Section III)'],
        ['Annual Premium:', '$1,284,000 (AIG layer only — see tower summary below)'],
        ['Producer:', 'Aon Global Risk Solutions — Technology Practice\n'
                     'Attn: Jennifer Walsh, Executive VP'],
    ]))
    story.append(Spacer(1, 10))

    # ── Tower ──
    story.append(Paragraph('SECTION I — LIABILITY TOWER STRUCTURE', styles['SectionHead']))
    story.append(Paragraph('The following table describes the complete liability tower for TechCorp Global. '
                           'AIG provides Layer 1 ($25M xs $5M). Higher layers are placed with separate carriers '
                           'as scheduled below. AIG has no obligations beyond the stated limit of Layer 1.',
                           styles['BodyText2']))
    story.append(Spacer(1, 6))

    story.append(styled_table(
        ['Layer', 'Carrier', 'Limit', 'Attachment\nPoint', 'Premium', 'Form', 'Defense'],
        [
            ['Layer 4', 'Berkshire Hathaway', '$50,000,000', '$80M xs\nunderlying', '$142,000', 'Follow-Form\n(AIG Lead)', 'Within\nLimits'],
            ['Layer 3', 'Swiss Re Corporate', '$25,000,000', '$55M xs\nunderlying', '$186,000', 'Follow-Form\n(AIG Lead)', 'Within\nLimits'],
            ['Layer 2', 'Chubb Excess', '$25,000,000', '$30M xs\nunderlying', '$324,000', 'Follow-Form\n(AIG Lead)', 'Within\nLimits'],
            ['Layer 1\n(AIG)', 'AIG Excess\nCasualty', '$25,000,000', '$5M xs\nunderlying', '$1,284,000', 'AIG-XS-2026\nLead Form', 'Outside\nLimits'],
            ['Underlying', 'Various\n(see §III)', '$5,000,000', 'Primary\n(see §III)', '$2,487,000', 'Various\nPrimary', 'Per\nUnderlying'],
        ],
        col_widths=[0.7 * inch, 1.0 * inch, 0.9 * inch, 0.8 * inch, 0.7 * inch, 0.9 * inch, 0.6 * inch]
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph('<b>Total Tower Limit:</b> $130,000,000  |  '
                           '<b>Total Program Premium:</b> $4,423,000  |  '
                           '<b>AIG Share:</b> 29% of total tower premium',
                           styles['BodyText2']))

    story.append(PageBreak())

    # ── Coverage ──
    story.append(Paragraph('SECTION II — COVERAGE PROVISIONS', styles['SectionHead']))
    story.append(Paragraph('<b>A. Insuring Agreement:</b> AIG will pay on behalf of the Insured those sums in '
                           'excess of the Retained Limit that the Insured becomes legally obligated to pay as '
                           'damages because of "bodily injury," "property damage," "personal and advertising injury," '
                           'or "products-completed operations hazard" to which this insurance applies.',
                           styles['BodyText2']))
    story.append(Spacer(1, 6))

    story.append(Paragraph('<b>B. Follow-Form Provisions:</b> Except as modified herein, this policy follows the '
                           'terms, conditions, definitions, and exclusions of the Scheduled Underlying Insurance. '
                           'Where the underlying policies contain conflicting terms, the most restrictive terms shall apply.',
                           styles['BodyText2']))
    story.append(Spacer(1, 6))

    story.append(Paragraph('<b>C. Drop-Down Provisions:</b>', styles['SubSection']))
    drops = [
        '1. <b>Exhaustion of Underlying Aggregate:</b> If any underlying aggregate limit is exhausted by payment '
        'of claims, this policy shall drop down and act as primary insurance for the remainder of the policy period, '
        'subject to the self-insured retention (SIR) of $50,000 per occurrence.',

        '2. <b>Underlying Insurer Insolvency:</b> If any underlying insurer becomes insolvent or is declared bankrupt, '
        'this policy shall not drop down. The Insured shall be responsible for the full amount of the underlying limit '
        'as if the underlying carrier had made payment.',

        '3. <b>Non-Renewal of Underlying:</b> If any underlying policy is cancelled or non-renewed during the term '
        'of this policy, the Insured must notify AIG within 30 days and secure replacement coverage acceptable to AIG '
        'within 60 days. Failure to do so will result in suspension of drop-down coverage.',

        '4. <b>Defense Costs:</b> Defense costs under Layer 1 (AIG) are payable OUTSIDE the limits of liability. '
        'AIG has the right and duty to defend any claim or suit seeking damages covered under this policy, even if '
        'the allegations are groundless, false, or fraudulent. Defense obligation ceases upon exhaustion of the '
        'applicable limit of liability.',

        '5. <b>Punitive/Exemplary Damages:</b> This policy provides coverage for punitive or exemplary damages '
        'where insurable by law in the jurisdiction where such damages are awarded. This coverage applies to the '
        'following jurisdictions only: CA, FL, GA, IL, NJ, NY, TX. All other jurisdictions: EXCLUDED.',
    ]
    for d in drops:
        story.append(Paragraph(d, styles['BodyText2']))
        story.append(Spacer(1, 3))

    story.append(PageBreak())

    # ── Underlying Schedule ──
    story.append(Paragraph('SECTION III — SCHEDULE OF UNDERLYING INSURANCE', styles['SectionHead']))
    story.append(Paragraph('The following policies must be maintained in force for the duration of this excess policy. '
                           'Any material change, cancellation, or non-renewal must be reported to AIG within 30 days.',
                           styles['BodyText2']))
    story.append(Spacer(1, 6))

    story.append(styled_table(
        ['Line', 'Carrier', 'Policy Number', 'Limit', 'Premium', 'Inception', 'Expiration'],
        [
            ['CGL', 'Hartford', 'HTF-GL-2026-44128', '$2M occ /\n$4M agg', '$486,000', '06/01/2026', '06/01/2027'],
            ['Products/\nComp Ops', 'Hartford', 'HTF-GL-2026-44128', '$2M occ /\n$4M agg', 'Included\nin CGL', '06/01/2026', '06/01/2027'],
            ['Commercial\nAuto', 'Travelers', 'TRV-CA-2026-88214', '$2M CSL', '$342,000', '06/01/2026', '06/01/2027'],
            ['Employers\'\nLiability', 'AIG', 'AIG-WC-2026-44283', '$1M ea acc /\n$1M disease', '$624,000', '06/01/2026', '06/01/2027'],
            ['D&O', 'Chubb', 'CHB-DO-2026-12847', '$5M per claim /\n$5M agg', '$428,000', '06/01/2026', '06/01/2027'],
            ['E&O / Tech\nProfessional', 'Beazley', 'BZL-EO-2026-34821', '$5M per claim /\n$5M agg', '$384,000', '06/01/2026', '06/01/2027'],
            ['Cyber\nLiability', 'AIG\nCyberEdge', 'AIG-CY-2026-22184', '$5M per event /\n$10M agg', '$223,000', '06/01/2026', '06/01/2027'],
        ],
        col_widths=[0.7 * inch, 0.7 * inch, 1.2 * inch, 0.9 * inch, 0.7 * inch, 0.7 * inch, 0.7 * inch]
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph('<b>Total Underlying Premium:</b> $2,487,000  |  '
                           '<b>Note:</b> The CGL and Products/Completed Operations share a common aggregate.',
                           styles['BodyText2']))
    story.append(Spacer(1, 10))

    # ── Exclusions ──
    story.append(Paragraph('SECTION IV — SPECIFIC EXCLUSIONS (Beyond Underlying)', styles['SectionHead']))
    exclusions = [
        '<b>1. AI / Autonomous Systems Exclusion:</b> This policy does not cover claims arising from or related to '
        'the operation, deployment, or failure of autonomous systems, machine learning models, or artificial intelligence '
        'applications developed, sold, or licensed by TechCorp Autonomous Systems, LLC, EXCEPT to the extent such claims '
        'are covered under the scheduled E&O/Technology Professional Liability policy (Beazley BZL-EO-2026-34821) and '
        'exhaust the underlying limit. This exception does not apply to autonomous vehicle or drone operations.',

        '<b>2. War / Terrorism / Cyber War:</b> This policy excludes all claims arising from: (a) war, invasion, '
        'hostilities, or warlike operations; (b) terrorism as defined under TRIA; (c) cyber warfare, state-sponsored '
        'cyber attacks, or attacks on critical infrastructure. Note: Standard cyber incidents are covered under the '
        'scheduled CyberEdge policy.',

        '<b>3. PFAS / Forever Chemicals:</b> Bodily injury or property damage arising from or related to the '
        'manufacture, distribution, sale, or use of per- and polyfluoroalkyl substances (PFAS) is EXCLUDED from '
        'this policy in its entirety. No coverage is provided regardless of the underlying policy terms.',

        '<b>4. Cannabis / THC Operations:</b> No coverage for claims arising from the cultivation, manufacture, '
        'distribution, sale, or use of cannabis or cannabis-derived products.',

        '<b>5. Sexual Abuse or Molestation:</b> This policy provides a $1,000,000 sublimit for sexual abuse or '
        'molestation claims (occurrence and aggregate). Defense costs are within this sublimit.',

        '<b>6. Employment Practices:</b> Employment practices claims (discrimination, harassment, wrongful termination, '
        'retaliation) are excluded unless covered under a scheduled underlying EPLI policy. TechCorp currently has no '
        'underlying EPLI — this gap is noted in the underwriting file.',
    ]
    for ex in exclusions:
        story.append(Paragraph(ex, styles['BodyText2']))
        story.append(Spacer(1, 4))

    story.append(PageBreak())

    # ── Loss History ──
    story.append(Paragraph('SECTION V — EXCESS LOSS HISTORY', styles['SectionHead']))
    story.append(Paragraph('Claims that have pierced or may pierce the underlying limits:', styles['BodyText2']))
    story.append(Spacer(1, 6))

    story.append(styled_table(
        ['Year', 'Claim', 'Line', 'Description', 'Underlying\nPaid', 'Excess\nPaid', 'Excess\nReserve', 'Status'],
        [
            ['2025', 'Patent\nInfringement', 'E&O', 'TechCorp Cloud v. Nexus Data — Patent '
             'infringement re: distributed caching algorithm. Trial verdict: $8.2M. Appeal pending.',
             '$5,000,000', '$1,800,000', '$1,400,000', 'OPEN\nAppeal'],
            ['2024', 'Data\nBreach', 'Cyber', '2.1M customer records exfiltrated via zero-day '
             'exploit. Notification costs, credit monitoring, regulatory fines (CCPA, GDPR). '
             'Class action settled.', '$5,000,000', '$3,200,000', '$0', 'CLOSED'],
            ['2023', 'Product\nLiability', 'CGL/\nProd', 'Autonomous drone collision at '
             'construction site. Worker suffered TBI. Suit alleged defective obstacle detection. '
             'Settled at mediation.', '$2,000,000', '$1,400,000', '$0', 'CLOSED'],
            ['2023', 'Securities\nClass Action', 'D&O', 'Shareholder suit re: revenue recognition '
             'restatement (FY2022 Q3-Q4). Dismissed on motion — defense costs only.',
             '$2,800,000', '$0', '$0', 'CLOSED\n(Defense)'],
        ],
        col_widths=[0.4 * inch, 0.6 * inch, 0.5 * inch, 2.0 * inch, 0.7 * inch, 0.6 * inch, 0.6 * inch, 0.6 * inch]
    ))
    story.append(Spacer(1, 8))
    story.append(Paragraph('<b>3-Year Excess Losses:</b> $7,800,000  |  <b>Outstanding Reserves:</b> $1,400,000  |  '
                           '<b>Total Excess Incurred:</b> $9,200,000',
                           styles['BodyText2']))
    story.append(Paragraph('<b>3-Year Layer 1 Loss Ratio:</b> 238% ($9.2M incurred vs. $3.85M premium over 3 years). '
                           'Renewal premium increase of 42% reflects this adverse experience.',
                           styles['BodyText2']))

    story.append(Spacer(1, 14))
    story.append(Paragraph('SECTION VI — SPECIAL UNDERWRITING CONDITIONS', styles['SectionHead']))
    conditions = [
        '<b>1. EPLI Gap:</b> TechCorp must procure Employment Practices Liability Insurance with a minimum limit '
        'of $5,000,000 by 09/01/2026. Failure to do so will result in a $250,000 sublimit for employment-related '
        'claims under this excess policy. TechCorp has 3,847 employees across 6 countries — EPLI is strongly recommended.',

        '<b>2. Autonomous Systems Review:</b> AIG Risk Engineering will conduct a comprehensive review of TechCorp '
        'Autonomous Systems\' safety protocols, testing procedures, and deployment guidelines by 12/31/2026. Findings '
        'may result in endorsement modifications at the next renewal.',

        '<b>3. International Coverage:</b> This policy provides worldwide coverage EXCEPT: Russia, Belarus, Iran, '
        'North Korea, Syria, Cuba, Venezuela, and Myanmar. Operations in the UK and Singapore are covered under the '
        'scheduled underlying policies which include DIC/DIL provisions for local admitted requirements.',
    ]
    for c in conditions:
        story.append(Paragraph(c, styles['BodyText2']))
        story.append(Spacer(1, 4))

    story.append(Spacer(1, 14))
    story.append(HRFlowable(width="100%", thickness=1, color=BORDER_GRAY))
    story.append(Spacer(1, 6))
    story.append(Paragraph('End of Policy Declaration — AIG Excess Casualty Division.', styles['SmallGray']))

    doc.build(story, onFirstPage=add_footer, onLaterPages=add_footer)
    print(f"  Created: {path}")


# ============================================================
# 3. CATASTROPHE LOSS REPORT — HURRICANE DAMAGE ASSESSMENT
#    (DOCX — detailed multi-section report)
# ============================================================
def generate_cat_loss_report():
    path = os.path.join(OUT_DIR, 'CAT_Loss_Report_Hurricane_Damage_CrescentBay.docx')
    doc = Document()

    title = doc.add_heading('Catastrophe Loss Report', level=0)
    for run in title.runs:
        run.font.color.rgb = AIG_BLUE_RGB

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        'Hurricane Helene — Category 3 Landfall\n'
        'CAT Event Reference: AIG-CAT-2025-HELENE-09\n'
        'Insured: Crescent Bay Resort & Marina, LLC\n'
        'Date of Loss: September 12-14, 2025\n'
        'Report Date: February 1, 2026  |  Valuation: January 15, 2026'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED_RGB

    doc.add_paragraph()

    # ── Executive Summary ──
    add_heading_styled(doc, 'Executive Summary', level=2)
    add_body(doc, 'Hurricane Helene made landfall as a Category 3 storm near Cedar Key, Florida on September 12, '
             '2025, with maximum sustained winds of 120 mph and a storm surge of 9-12 feet above normal tide levels '
             'along the Gulf Coast of Florida. The insured property, Crescent Bay Resort & Marina, sustained severe '
             'wind damage, storm surge flooding, and marina infrastructure destruction.')
    add_body(doc, 'TOTAL ESTIMATED LOSS: $18,742,000 (subject to adjustment)', bold=True)
    add_body(doc, 'POLICY LIMITS ADEQUACY: Estimated loss exceeds property limit. BI/EE approaching sublimit. '
             'Excess carrier (Chubb) notified on 09/15/2025. Reinsurance notification filed under Treaty XL-2025-004.', bold=True)

    doc.add_paragraph()

    # ── Policy Information ──
    add_heading_styled(doc, 'Policy Information', level=2)
    policy_table = doc.add_table(rows=14, cols=2, style='Table Grid')
    policy_data = [
        ('Policy Number', 'AIG-CPP-2025-FL-71284'),
        ('Named Insured', 'Crescent Bay Resort & Marina, LLC'),
        ('Additional Insured', 'Crescent Bay Holdings, Inc. (Parent)\nFirst National Bank of Tampa (Mortgagee)'),
        ('Property Address', '4800 Gulf Shore Drive, Crystal River, FL 34429'),
        ('Building Limit', '$12,500,000 (Replacement Cost)'),
        ('BPP / Contents Limit', '$4,200,000 (Replacement Cost)'),
        ('Business Income + EE', '$3,500,000 (12-month ALS, 72-hr waiting period)'),
        ('Named Storm Deductible', '5% of TIV at time of loss = $1,010,000 (based on $20,200,000 TIV)'),
        ('Flood Sublimit', '$2,500,000 per occurrence (Zone VE — High Risk Coastal)'),
        ('Flood Deductible', '$500,000'),
        ('Marina/Pier Sublimit', '$1,500,000 (ACV basis)'),
        ('Ordinance or Law', '$1,250,000 (Coverages A, B, C combined)'),
        ('Debris Removal', '$500,000 (additional)'),
        ('Policy Period', '07/01/2025 — 07/01/2026'),
    ]
    for i, (k, v) in enumerate(policy_data):
        policy_table.cell(i, 0).text = k
        policy_table.cell(i, 1).text = v
        for p in policy_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in policy_table.cell(i, 1).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_page_break()

    # ── Storm Timeline ──
    add_heading_styled(doc, 'Storm Event Timeline', level=2)
    timeline = [
        ('09/09/2025 — 72 hrs prior', 'NHC issues Tropical Storm Watch for Citrus County. Insured activates '
         'hurricane preparedness plan. Marina vessels begin evacuation. Portable generators staged.'),
        ('09/10/2025 — 48 hrs prior', 'Tropical Storm Watch upgraded to Hurricane Warning. Helene intensifies to '
         'Category 2. Insured deploys storm shutters, sandbags lobby and ground-floor rooms. Pool furniture secured. '
         'Guests evacuated per county mandatory evacuation order (Zone A).'),
        ('09/11/2025 — 24 hrs prior', 'Helene rapidly intensifies to Category 3. Storm surge forecast increased '
         'from 6-8 ft to 9-12 ft. All staff except essential personnel evacuated. Emergency power systems tested. '
         'Final property walkthrough conducted and documented with video.'),
        ('09/12/2025 — Landfall', 'Hurricane Helene makes landfall at 2:14 AM EDT near Cedar Key (35 miles south '
         'of insured property) with 120 mph sustained winds. Property experiences 105 mph sustained winds with '
         '128 mph gusts. Storm surge reaches 10.4 ft above normal at property location per USGS gauge #02310947.'),
        ('09/12/2025 — Post-Landfall', 'Storm surge begins receding by 8:00 AM. Significant structural damage '
         'visible. Marina pier partially collapsed. 4 of 6 buildings show roof damage. Ground floor flooded to '
         '4.2 ft depth. Emergency services inaccessible until 09/13.'),
        ('09/13/2025 — First Assessment', 'Insured conducts initial damage assessment with property manager and '
         'maintenance team. Photos and video documentation compiled. AIG notified via FNOL at 10:42 AM. '
         'Claim number assigned: CLM-2025-FL-CAT-42871.'),
        ('09/15/2025 — Adjuster Deployed', 'AIG CAT adjuster (Thomas Brennan, Senior Complex Claims) arrives on-site. '
         'Independent structural engineer (Wiss, Janney, Elstner Associates) retained. Forensic meteorologist '
         '(Haag Engineering) retained for wind vs. water causation analysis.'),
    ]
    for date, desc in timeline:
        add_body(doc, date, bold=True)
        add_body(doc, desc)

    doc.add_page_break()

    # ── Damage Assessment ──
    add_heading_styled(doc, 'Detailed Damage Assessment', level=2)

    add_heading_styled(doc, 'Building 1: Main Resort (3-story, 84 rooms)', level=3)
    add_body(doc, 'Construction: Reinforced concrete frame, stucco exterior, concrete tile roof (installed 2018).')
    b1 = [
        ('Roof System', 'Wind', '$1,420,000', 'Approximately 35% of concrete tiles displaced or broken. '
         'Underlayment exposed in 6 areas. Water intrusion into 22 guest rooms on 3rd floor. '
         'Tile manufacturer (Eagle Roofing) confirmed tiles rated for 150 mph — failure at 128 mph '
         'gusts suggests installation deficiency, not product failure. Roof warranty claim denied by installer.'),
        ('Building Envelope', 'Wind', '$680,000', 'Impact damage to stucco facade on east and south elevations. '
         '14 windows broken (non-impact rated — pre-dates FL Building Code 2020 requirement). '
         'Balcony railings damaged on 8 units.'),
        ('Interior — 3rd Floor', 'Wind/\nWater', '$1,240,000', 'Water damage to 22 guest rooms: flooring, drywall, '
         'furniture, fixtures. Mold remediation required in 16 rooms (testing confirmed Aspergillus/Penicillium '
         'above acceptable levels). Full gut renovation required.'),
        ('Interior — 2nd Floor', 'Wind-\nDriven\nRain', '$420,000', 'Wind-driven rain penetration through window '
         'breaches. 8 rooms affected — carpeting, lower drywall, bathroom fixtures.'),
        ('Interior — 1st Floor', 'Storm\nSurge', '$1,860,000', 'Storm surge flooding to 4.2 ft depth. Complete '
         'destruction of lobby, restaurant, spa, fitness center, and 12 ground-floor rooms. All flooring, '
         'drywall to 4 ft, electrical outlets/switches, kitchen equipment, and HVAC air handlers destroyed. '
         'NOTE: Subject to Flood Sublimit and separate deductible.'),
        ('Elevator Systems', 'Flood', '$340,000', 'Both elevator pits flooded. Control panels, motors, and '
         'guide rails require replacement. Otis Elevator inspection confirms total loss of elevator machinery. '
         'Estimated replacement: 16-20 weeks lead time.'),
        ('Electrical Systems', 'Surge/\nFlood', '$520,000', 'Main electrical panel (ground floor) destroyed by '
         'surge water. Emergency generator operated for 72 hours post-storm but sustained salt water damage to '
         'cooling system. Transfer switch damaged. Complete rewiring of ground floor required.'),
    ]

    b1_table = doc.add_table(rows=len(b1) + 1, cols=4, style='Table Grid')
    b1_headers = ['Component', 'Peril', 'Estimated\nCost', 'Description']
    for j, h in enumerate(b1_headers):
        b1_table.cell(0, j).text = h
    style_header_row(b1_table.rows[0])
    for i, (comp, peril, cost, desc) in enumerate(b1):
        b1_table.cell(i + 1, 0).text = comp
        b1_table.cell(i + 1, 1).text = peril
        b1_table.cell(i + 1, 2).text = cost
        b1_table.cell(i + 1, 3).text = desc
        for j in range(4):
            for p in b1_table.cell(i + 1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    add_body(doc, 'Building 1 Subtotal: $6,480,000', bold=True)

    doc.add_page_break()

    add_heading_styled(doc, 'Building 2: Conference Center & Ballroom', level=3)
    add_body(doc, 'Construction: Steel frame, glass curtain wall (south), metal roof. Built 2015.')
    b2_items = [
        ('Metal Roof', 'Wind', '$480,000', 'Standing seam metal roof — 20% of panels peeled back by wind uplift. '
         'Ridge cap separation along 120 linear feet. Underlayment intact in most areas.'),
        ('Glass Curtain Wall', 'Wind/\nDebris', '$890,000', 'South-facing glass curtain wall suffered catastrophic '
         'failure. 8 of 24 panels shattered by windborne debris (tree limbs from neighboring property). Interior '
         'exposed to wind-driven rain for estimated 6 hours before tarping. AV equipment, stage, and ballroom '
         'flooring destroyed.'),
        ('Interior', 'Water', '$640,000', 'Water damage throughout from glass failure and roof leaks. Drop ceiling '
         'collapsed. Carpet, lighting, and wall finishes require full replacement. Kitchen (catering) equipment '
         'damaged by water intrusion.'),
        ('HVAC', 'Wind', '$180,000', '3 rooftop HVAC units displaced from curbs. Ductwork damaged. Refrigerant '
         'released (R-410A — EPA notification required).'),
    ]
    for comp, peril, cost, desc in b2_items:
        add_body(doc, f'{comp} ({peril}): {cost}', bold=True)
        add_body(doc, desc)
    add_body(doc, 'Building 2 Subtotal: $2,190,000', bold=True)

    doc.add_paragraph()

    add_heading_styled(doc, 'Marina & Pier Infrastructure', level=3)
    marina_items = [
        ('Fixed Pier (300 ft)', 'Storm\nSurge', '$680,000', 'Concrete pier deck separated from pilings at 3 points. '
         'Approximately 120 ft of pier is unusable. Pilings intact below waterline per dive inspection.'),
        ('Floating Dock System', 'Surge/\nWind', '$420,000', 'Floating dock system (40 slips) — 60% destroyed. '
         'Dock fingers broken, cleats pulled, gangways twisted. 12 slips salvageable with repair.'),
        ('Fuel Dock & Pump Station', 'Surge', '$340,000', 'Underground fuel tanks intact (double-walled FRP). '
         'Above-ground dispensers, piping, and electrical destroyed. DEP inspection required before reconstruction. '
         'Fuel inventory loss: 2,400 gallons diesel, 1,800 gallons gasoline ($18,200).'),
        ('Ship Store & Marina Office', 'Surge', '$180,000', 'Single-story CBS structure — flooded to 3.8 ft. '
         'Contents total loss. Structure repairable.'),
        ('Boat Lift (2 units)', 'Wind', '$120,000', '1 of 2 boat lifts twisted off rails. Motor housing damaged '
         'on both units. Replacement parts: 12-16 week lead time from manufacturer.'),
    ]
    for comp, peril, cost, desc in marina_items:
        add_body(doc, f'{comp} ({peril}): {cost}', bold=True)
        add_body(doc, desc)
    add_body(doc, 'Marina Subtotal: $1,740,000 (subject to $1,500,000 sublimit — excess of $240,000 not covered)',
             bold=True)

    doc.add_page_break()

    # ── Coverage Analysis ──
    add_heading_styled(doc, 'Coverage Application & Loss Summary', level=2)

    loss_summary = doc.add_table(rows=14, cols=4, style='Table Grid')
    ls_headers = ['Category', 'Gross Loss', 'Deductible /\nSublimit', 'Net Payable']
    for j, h in enumerate(ls_headers):
        loss_summary.cell(0, j).text = h
    style_header_row(loss_summary.rows[0])

    ls_data = [
        ('Building 1 — Wind Damage', '$3,760,000', 'Named Storm\nDed: $1,010,000', '$2,750,000'),
        ('Building 1 — Storm Surge (Flood)', '$2,720,000', 'Flood Sub:\n$2,500,000\nFlood Ded: $500,000', '$2,000,000'),
        ('Building 2 — Wind Damage', '$1,550,000', 'Named Storm\nDed (applied\nabove)', '$1,550,000'),
        ('Building 2 — Water Damage', '$640,000', '—', '$640,000'),
        ('Marina & Pier', '$1,740,000', 'Sublimit:\n$1,500,000', '$1,500,000'),
        ('Other Buildings (3-6)', '$1,420,000', '—', '$1,420,000'),
        ('BPP / Contents (all bldgs)', '$2,180,000', '—', '$2,180,000'),
        ('Debris Removal', '$482,000', 'Add\'l limit:\n$500,000', '$482,000'),
        ('Business Income (5.5 months)', '$2,640,000', '72-hr wait\napplied', '$2,640,000'),
        ('Extra Expense', '$860,000', 'Combined\nw/ BI sub:\n$3,500,000', '$860,000'),
        ('Ordinance or Law', '$550,000', 'Sublimit:\n$1,250,000', '$550,000'),
        ('Emergency Board-Up / Tarping', '$148,000', 'Sue & Labor', '$148,000'),
        ('TOTAL', '$18,690,000', 'Various', '$16,720,000'),
    ]
    for i, (cat, gross, ded, net) in enumerate(ls_data):
        loss_summary.cell(i + 1, 0).text = cat
        loss_summary.cell(i + 1, 1).text = gross
        loss_summary.cell(i + 1, 2).text = ded
        loss_summary.cell(i + 1, 3).text = net
        for j in range(4):
            for p in loss_summary.cell(i + 1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    if i == len(ls_data) - 1:
                        run.bold = True

    doc.add_paragraph()
    add_body(doc, 'COVERAGE GAP ANALYSIS:', bold=True)
    gaps = [
        'Marina excess above sublimit: $240,000 — NOT COVERED (recommend sublimit increase at renewal)',
        'Named Storm deductible impact: $1,010,000 — insured bears this cost',
        'Flood deductible: $500,000 — insured bears this cost',
        'Business Income approaching sublimit: $3,500,000 (current claim: $3,500,000) — if restoration exceeds '
        '5.5 months, additional BI losses will be UNINSURED',
        'Code upgrade costs (O&L): $550,000 estimated, well within $1,250,000 sublimit — ADEQUATELY COVERED',
        'Total building damage ($12,970,000) approaches building limit ($12,500,000) — potential $470,000 shortfall '
        'if additional damage discovered during reconstruction',
    ]
    for g in gaps:
        add_body(doc, f'  - {g}')

    doc.add_page_break()

    # ── Wind vs Water ──
    add_heading_styled(doc, 'Wind vs. Water Causation Analysis', level=2)
    add_body(doc, 'Prepared by: Haag Engineering Company — Forensic Meteorology Division')
    add_body(doc, 'The causation of damage at Crescent Bay Resort requires careful allocation between wind (covered '
             'under Named Storm provisions) and water/flood (subject to separate sublimit and deductible). Haag '
             'Engineering conducted on-site inspection on 09/18-19/2025 and reviewed the following data sources:')
    sources = [
        'USGS storm surge gauge data (#02310947) — peak surge 10.4 ft NAVD88',
        'NOAA buoy data (Station 42036) — peak sustained wind 105 mph, gusts 128 mph',
        'Florida Automated Weather Network (FAWN) station at Lecanto — peak gust 118 mph',
        'ASCE 7-22 wind speed analysis for Risk Category III structures',
        'Insured\'s pre-storm video documentation (time-stamped 09/11/2025)',
        'Satellite imagery (Maxar/Planet Labs) — pre-storm (09/10) and post-storm (09/13)',
    ]
    for s in sources:
        add_body(doc, f'  {s}')

    add_body(doc, '\nFINDINGS:', bold=True)
    add_body(doc, 'Based on high-water marks, debris field analysis, structural damage patterns, and the timeline '
             'of storm conditions, Haag Engineering concludes the following causation allocation:')

    alloc_table = doc.add_table(rows=8, cols=3, style='Table Grid')
    alloc_headers = ['Building / Component', 'Wind %', 'Water %']
    for j, h in enumerate(alloc_headers):
        alloc_table.cell(0, j).text = h
    style_header_row(alloc_table.rows[0])
    alloc_data = [
        ('Building 1 — Roof & Upper Floors', '95%', '5%'),
        ('Building 1 — Ground Floor', '15%', '85%'),
        ('Building 2 — Conference Center', '80%', '20%'),
        ('Marina & Pier', '25%', '75%'),
        ('Other Buildings (above 1st floor)', '90%', '10%'),
        ('Other Buildings (ground floor)', '20%', '80%'),
        ('Site Improvements & Landscaping', '40%', '60%'),
    ]
    for i, (comp, wind, water) in enumerate(alloc_data):
        alloc_table.cell(i + 1, 0).text = comp
        alloc_table.cell(i + 1, 1).text = wind
        alloc_table.cell(i + 1, 2).text = water
        for j in range(3):
            for p in alloc_table.cell(i + 1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    add_body(doc, 'NOTE: The anti-concurrent causation clause in the Named Storm endorsement provides that where '
             'wind and water damage cannot be separated, the loss is subject to the higher deductible. However, '
             'Haag Engineering\'s analysis provides a defensible allocation that AIG recommends adopting for '
             'claim adjustment purposes, subject to insured agreement.')

    doc.add_paragraph()

    # ── Next Steps ──
    add_heading_styled(doc, 'Adjuster Recommendations & Next Steps', level=2)
    steps = [
        '1. Complete structural engineering report (WJE) — expected by 02/28/2026',
        '2. Finalize wind/water causation allocation with insured and mortgagee',
        '3. Issue advance payment of $5,000,000 (undisputed wind damage) pending final adjustment',
        '4. Engage replacement cost estimator for Building 1 renovation scope',
        '5. Monitor Business Income claim — insured projects reopening 03/15/2026 (partial) and 06/01/2026 (full)',
        '6. Coordinate with excess carrier (Chubb) if building damage exceeds primary limit',
        '7. Review Ordinance or Law requirements with Citrus County Building Department',
        '8. File subrogation potential against roof tile installer (warranty/installation deficiency)',
        '9. DEP clearance required before marina fuel dock reconstruction',
        '10. Reinsurance bordereaux update — Treaty XL-2025-004, AIG retention: $5,000,000',
    ]
    for s in steps:
        add_body(doc, s)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Report prepared by: Thomas Brennan, Senior Complex Claims Adjuster, AIG CAT Response Team. '
                    'This report is confidential and prepared for internal claim management purposes only. '
                    'All figures are preliminary estimates subject to final adjustment.')
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED_RGB
    run.italic = True

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# 4. D&O + FIDUCIARY LIABILITY — COMPLEX CLAIMS-MADE POLICY
#    (DOCX — dense policy with sublimits, retentions, exclusions)
# ============================================================
def generate_do_fiduciary():
    path = os.path.join(OUT_DIR, 'DO_Fiduciary_Liability_Pinnacle_Healthcare.docx')
    doc = Document()

    title = doc.add_heading('Directors & Officers + Fiduciary Liability', level=0)
    for run in title.runs:
        run.font.color.rgb = AIG_BLUE_RGB

    subtitle = doc.add_paragraph()
    run = subtitle.add_run(
        'Claims-Made and Reported Policy\n'
        'Policy No: AIG-DO-2026-PH-55821  |  Fiduciary: AIG-FID-2026-PH-55822\n'
        'Pinnacle Healthcare Systems, Inc.\n'
        'Effective: 01/01/2026 — 01/01/2027'
    )
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED_RGB

    doc.add_paragraph()

    # ── Declarations ──
    add_heading_styled(doc, 'Part I — Declarations', level=2)
    decl_table = doc.add_table(rows=16, cols=2, style='Table Grid')
    decl_data = [
        ('Named Insured (Organization)', 'Pinnacle Healthcare Systems, Inc.'),
        ('Subsidiaries', 'Pinnacle Medical Group, P.A.\n'
                        'Pinnacle Surgery Centers, LLC (12 locations)\n'
                        'Pinnacle Home Health Services, Inc.\n'
                        'Pinnacle Health IT Solutions, LLC\n'
                        'Pinnacle Physician Staffing, Inc.\n'
                        'Pinnacle Healthcare Foundation (501(c)(3))'),
        ('State of Incorporation', 'Delaware'),
        ('Publicly Traded?', 'Yes — NYSE: PNHL (Market Cap: $4.2B as of 12/31/2025)'),
        ('Annual Revenue', '$3.84 Billion (FY 2025)'),
        ('Total Assets', '$6.12 Billion'),
        ('Total Employees', '18,420 (including 2,840 physicians)'),
        ('Board of Directors', '11 members (8 independent, 3 management)'),
        ('D&O Limit of Liability', '$15,000,000 per claim and annual aggregate\n'
                                   '(Shared across Side A, B, and C)'),
        ('Side A (Non-Indemnifiable)', '$15,000,000 — No retention'),
        ('Side B (Corporate Reimbursement)', '$15,000,000 — $500,000 corporate retention per claim'),
        ('Side C (Entity Securities)', '$15,000,000 — $1,000,000 corporate retention per claim\n'
                                       'Sublimit: $7,500,000 for entity-only securities claims'),
        ('Fiduciary Liability Limit', '$5,000,000 per claim and aggregate\n'
                                      '$25,000 retention per claim'),
        ('Retroactive Date', 'January 1, 2018 (D&O)  |  January 1, 2020 (Fiduciary)'),
        ('Extended Reporting Period', 'Automatic: 60 days  |  Optional: 1-year ($100% premium)\n'
                                      '3-year ($200% premium)'),
        ('Annual Premium', 'D&O: $842,000  |  Fiduciary: $124,000  |  Total: $966,000'),
    ]
    for i, (k, v) in enumerate(decl_data):
        decl_table.cell(i, 0).text = k
        decl_table.cell(i, 1).text = v
        for p in decl_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in decl_table.cell(i, 1).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_page_break()

    # ── D&O Tower ──
    add_heading_styled(doc, 'D&O Liability Tower', level=2)
    add_body(doc, 'The following tower provides $50,000,000 in total D&O limits. AIG is the primary carrier (Layer 1).')

    tower = doc.add_table(rows=5, cols=5, style='Table Grid')
    tower_headers = ['Layer', 'Carrier', 'Limit', 'Attachment', 'Premium']
    for j, h in enumerate(tower_headers):
        tower.cell(0, j).text = h
    style_header_row(tower.rows[0])
    tower_data = [
        ('Layer 4', 'QBE', '$10,000,000', 'xs $40M', '$86,000'),
        ('Layer 3', 'Zurich', '$15,000,000', 'xs $25M', '$148,000'),
        ('Layer 2', 'Chubb', '$10,000,000', 'xs $15M', '$224,000'),
        ('Layer 1 (AIG)', 'AIG', '$15,000,000', 'Primary', '$842,000'),
    ]
    for i, (layer, carrier, limit, attach, prem) in enumerate(tower_data):
        tower.cell(i + 1, 0).text = layer
        tower.cell(i + 1, 1).text = carrier
        tower.cell(i + 1, 2).text = limit
        tower.cell(i + 1, 3).text = attach
        tower.cell(i + 1, 4).text = prem
        for j in range(5):
            for p in tower.cell(i + 1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    add_body(doc, 'Total D&O Tower: $50,000,000  |  Total Tower Premium: $1,300,000', bold=True)

    doc.add_paragraph()

    # ── Insured Persons ──
    add_heading_styled(doc, 'Part II — Insured Persons & Coverage', level=2)
    add_body(doc, '"Insured Person" means any natural person who was, is, or becomes during the Policy Period:')
    persons = [
        '(a) A director, officer, or equivalent executive of the Named Insured or any Subsidiary',
        '(b) A trustee, administrator, or fiduciary of any Employee Benefit Plan',
        '(c) A shadow director or de facto director as determined by applicable law',
        '(d) An employee of the Named Insured ONLY with respect to Claims alleging wrongful '
        'employment practices (EPLI extension — see Part V)',
        '(e) The estate, heirs, or legal representatives of a deceased Insured Person',
        '(f) The lawful spouse or domestic partner of an Insured Person ONLY to the extent '
        'Claims seek damages from marital community property or jointly held assets',
    ]
    for p_text in persons:
        add_body(doc, p_text)

    doc.add_paragraph()

    # ── Coverage Sections ──
    add_heading_styled(doc, 'Part III — Coverage Sections (D&O)', level=2)

    sections = [
        ('Side A — Non-Indemnifiable Loss',
         'AIG shall pay Loss on behalf of any Insured Person arising from a Claim for a Wrongful Act, '
         'but only to the extent such Loss is not indemnified by the Organization and is not indemnifiable '
         'pursuant to applicable law, the Organization\'s charter, bylaws, or agreements.\n\n'
         'Side A applies with NO retention. This is the broadest coverage and protects individual directors '
         'and officers when the company cannot or will not indemnify them (e.g., insolvency, derivative suits, '
         'or legal prohibition on indemnification).\n\n'
         'Side A DIC (Difference in Conditions) Feature: If any underlying or co-primary D&O policy denies '
         'coverage for a claim that would be covered under this policy\'s terms, Side A DIC drops down to '
         'provide coverage, subject to the full policy limit.'),

        ('Side B — Corporate Reimbursement',
         'AIG shall pay Loss on behalf of the Organization arising from a Claim for a Wrongful Act committed '
         'by an Insured Person, but only to the extent the Organization has indemnified or is required to '
         'indemnify the Insured Person.\n\n'
         'Retention: $500,000 per Claim applies to Side B only. The retention is satisfied by the Organization\'s '
         'payment (not by AIG payment). If the Organization fails to advance the retention amount within 90 days '
         'of demand by the Insured Person, Side A coverage shall apply as if the loss were non-indemnifiable.'),

        ('Side C — Entity Securities Coverage',
         'AIG shall pay Loss on behalf of the Organization arising from a Securities Claim brought by a '
         'security holder of the Organization, whether individually, derivatively, or as part of a class, '
         'alleging a violation of any federal, state, or foreign securities law or regulation.\n\n'
         'Retention: $1,000,000 per Securities Claim.\n'
         'Sublimit: $7,500,000 for entity-only Securities Claims (where no individual Insured Person is named).\n\n'
         'IMPORTANT: Securities Claims arising from the following are subject to heightened scrutiny and may '
         'require notice to AIG within 15 days (instead of standard 60 days):\n'
         '- Restatement of financial statements\n'
         '- SEC Wells Notice or formal investigation\n'
         '- Whistleblower complaint under Dodd-Frank\n'
         '- Shareholder demand letter alleging breach of fiduciary duty\n'
         '- M&A transaction litigation (including appraisal actions)'),
    ]

    for heading, body in sections:
        add_heading_styled(doc, heading, level=3)
        add_body(doc, body)

    doc.add_page_break()

    # ── Exclusions ──
    add_heading_styled(doc, 'Part IV — Exclusions', level=2)
    add_body(doc, 'This policy shall not cover any Loss arising from or related to:')

    exclusions = [
        ('A. Prior & Pending Litigation', 'Any Claim based upon, arising from, or attributable to any '
         'litigation, administrative proceeding, or investigation pending on or before the Retroactive Date, '
         'or based upon the same or substantially the same facts alleged in such prior proceedings. '
         'Pending matters as of the Retroactive Date are listed in Schedule A attached.'),

        ('B. Fraud & Criminal Acts', 'Any deliberately fraudulent, dishonest, or criminal act committed by '
         'an Insured Person, as established by a final, non-appealable adjudication, admission, or plea. '
         'Defense costs are advanced until such determination is made. If fraud is established, AIG shall '
         'be entitled to full reimbursement of all defense costs advanced (severability preserved — see Part VII).'),

        ('C. Personal Profit / Illegal Remuneration', 'Any Claim based upon an Insured Person gaining '
         'personal profit, remuneration, or advantage to which they were not legally entitled, including '
         'but not limited to: insider trading, Section 16(b) short-swing profits, undisclosed self-dealing '
         'transactions, or undisclosed related-party transactions.'),

        ('D. Professional Medical Services', 'Any Claim arising from the rendering or failure to render '
         'professional medical, surgical, dental, or pharmaceutical services. This exclusion does NOT apply '
         'to: (1) Claims against directors or officers for failure to supervise the delivery of medical '
         'services; (2) Claims alleging negligent credentialing or privileging decisions; or (3) Claims '
         'alleging EMTALA violations at the board/management level.'),

        ('E. ERISA / Employee Benefits', 'Any Claim under ERISA (Employee Retirement Income Security Act) '
         'or similar state law relating to the administration of any Employee Benefit Plan, EXCEPT as '
         'covered under the Fiduciary Liability section (Part VI) of this policy.'),

        ('F. Antitrust / Price Fixing', 'Any Claim alleging violation of the Sherman Antitrust Act, '
         'Clayton Act, Robinson-Patman Act, or any state antitrust or unfair competition statute, '
         'EXCEPT defense costs only are covered until a final determination of liability. '
         'Sublimit for antitrust defense: $2,500,000.'),

        ('G. Pollution / Environmental', 'Any Claim for bodily injury, property damage, or remediation '
         'costs arising from the actual, alleged, or threatened discharge, dispersal, release, or escape '
         'of pollutants or contaminants. This exclusion does NOT apply to Securities Claims alleging '
         'failure to disclose environmental liabilities in SEC filings.'),

        ('H. Cyber / Privacy Event', 'Any Claim arising from a data breach, ransomware attack, or privacy '
         'violation, EXCEPT to the extent such Claim is brought as a Securities Claim alleging failure to '
         'maintain adequate cybersecurity controls or failure to disclose a material cyber incident. '
         'Note: Dedicated Cyber Liability coverage should be maintained separately.'),

        ('I. Medicare / Medicaid Fraud', 'Any Claim brought by or on behalf of the United States Government '
         'alleging violation of the False Claims Act (31 USC §3729), Anti-Kickback Statute (42 USC §1320a-7b), '
         'or Stark Law (42 USC §1395nn) relating to Medicare or Medicaid billing practices. '
         'IMPORTANT: This exclusion applies to entity coverage (Side C) only. Individual directors and '
         'officers retain coverage under Side A and B for defense of such claims.'),
    ]

    for heading, body in exclusions:
        add_heading_styled(doc, heading, level=3)
        add_body(doc, body)

    doc.add_page_break()

    # ── EPLI Extension ──
    add_heading_styled(doc, 'Part V — Employment Practices Liability Extension', level=2)
    add_body(doc, 'This policy includes an Employment Practices Liability (EPLI) coverage extension for '
             'Pinnacle Healthcare Systems with the following terms:')

    epli_table = doc.add_table(rows=8, cols=2, style='Table Grid')
    epli_data = [
        ('EPLI Sublimit', '$5,000,000 per claim and aggregate (within D&O limit)'),
        ('EPLI Retention', '$150,000 per claim (corporate)'),
        ('Covered Wrongful Acts', 'Discrimination, harassment (including sexual harassment), wrongful '
         'termination, retaliation, failure to promote, breach of employment contract, negligent '
         'evaluation, defamation in employment context, invasion of employee privacy'),
        ('Third-Party EPLI', 'Included — covers claims by patients, vendors, or visitors alleging '
         'harassment or discrimination by employees'),
        ('Wage & Hour Exclusion', 'EXCLUDED — No coverage for FLSA or state wage/hour claims. '
         'Pinnacle had a $4.2M FLSA class action settlement in 2023 involving nurse overtime.'),
        ('Workplace Violence', 'Included — $1,000,000 sublimit for workplace violence response costs '
         '(counseling, security, relocation)'),
        ('Regulatory Proceedings', 'EEOC, state human rights agencies, DOL investigations — '
         'defense costs covered (no damages coverage for fines/penalties)'),
        ('Predecessor Coverage', 'Includes claims arising from employment practices of acquired entities '
         'provided acquisition closed after the Retroactive Date (01/01/2018)'),
    ]
    for i, (k, v) in enumerate(epli_data):
        epli_table.cell(i, 0).text = k
        epli_table.cell(i, 1).text = v
        for p in epli_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in epli_table.cell(i, 1).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # ── Fiduciary ──
    add_heading_styled(doc, 'Part VI — Fiduciary Liability Coverage', level=2)
    add_body(doc, 'This section provides separate coverage for claims alleging breach of fiduciary duty in '
             'connection with the administration of Employee Benefit Plans maintained by Pinnacle Healthcare.')

    fid_plans = [
        'Pinnacle Healthcare 401(k) Retirement Savings Plan — $482M assets (Fidelity)',
        'Pinnacle Healthcare Defined Benefit Pension Plan — $218M assets (TIAA) — FROZEN effective 01/01/2022',
        'Pinnacle Healthcare Group Health Plan — self-insured, 18,420 participants (Aetna ASO)',
        'Pinnacle Healthcare Executive Deferred Compensation Plan — $34M (Rabbi Trust, Vanguard)',
        'Pinnacle Healthcare ESOP — 2.4M shares ($48M at current market)',
    ]
    for plan in fid_plans:
        add_body(doc, f'  {plan}')

    doc.add_paragraph()
    add_body(doc, 'FIDUCIARY RISK FACTORS:', bold=True)
    risks = [
        'The 401(k) plan was the subject of an excessive fee lawsuit (Garcia v. Pinnacle Healthcare, '
        'S.D. Tex. 2024) that was dismissed on summary judgment. Plaintiffs have filed a notice of appeal '
        'to the Fifth Circuit. If reversed, potential exposure is estimated at $8-12M in fee disgorgement. '
        'This matter is SCHEDULED as a Known Circumstance and is subject to a $2,000,000 sublimit.',

        'The frozen Defined Benefit Plan is currently 87% funded. PBGC premiums are current. The plan\'s '
        'investment committee (3 members of the Board) oversees asset allocation. Recent shift from fixed '
        'income to alternative investments (12% allocation to PE/hedge funds) may attract fiduciary scrutiny.',

        'The ESOP underwent a Section 409(a) valuation in December 2025. The independent appraiser '
        '(Stout Risius Ross) valued shares at $20.14 per share. Market price on valuation date was '
        '$17.42. The premium of 15.6% over market is within acceptable range per DOL guidance but could '
        'face challenge if stock price declines significantly.',
    ]
    for r in risks:
        add_body(doc, r)

    doc.add_page_break()

    # ── Loss History ──
    add_heading_styled(doc, 'Part VIII — Claims History (5 Years)', level=2)

    claims_table = doc.add_table(rows=8, cols=6, style='Table Grid')
    claims_headers = ['Year', 'Claim Type', 'Description', 'Paid', 'Reserve', 'Status']
    for j, h in enumerate(claims_headers):
        claims_table.cell(0, j).text = h
    style_header_row(claims_table.rows[0])

    claims_data = [
        ('2025', 'Securities\nClass Action', 'Shareholder suit alleging misleading statements '
         'about patient volume projections in Q2 2025 earnings call. Lead plaintiff: '
         'Midwest Pension Trust.', '$0', '$4,200,000', 'OPEN\nDisc.'),
        ('2024', 'EPLI', 'Former Chief Nursing Officer alleged wrongful termination '
         'and gender discrimination. Settled at mediation.', '$680,000', '$0', 'CLOSED'),
        ('2024', 'Derivative\nSuit', 'Shareholder derivative action alleging waste of '
         'corporate assets related to failed EHR system implementation ($42M write-off). '
         'Demand refused. Suit filed in Delaware Chancery.', '$1,240,000', '$0', 'CLOSED\n(Settled)'),
        ('2023', 'FLSA\nClass Action', 'Nurse overtime class action — 3,200 class members. '
         'NOT COVERED under D&O (wage/hour exclusion). Company paid from operating funds.',
         'N/A\n(Excluded)', 'N/A', 'CLOSED\n$4.2M'),
        ('2023', 'Fiduciary', '401(k) excessive fee suit (Garcia v. Pinnacle). Defense costs '
         'covered under fiduciary policy. Dismissed on SJ — appeal pending.',
         '$486,000\n(Defense)', '$200,000', 'OPEN\nAppeal'),
        ('2022', 'Regulatory', 'DOJ/OIG investigation re: billing practices at 3 surgery '
         'centers. No charges filed. Defense costs only.',
         '$1,840,000\n(Defense)', '$0', 'CLOSED'),
        ('2021', 'D&O\n(Side A)', 'Former CEO/founder sued individually by PE investor '
         'alleging breach of representations in 2019 recapitalization. Non-indemnifiable '
         'by company per Delaware law. Settled confidentially.',
         '$2,400,000', '$0', 'CLOSED'),
    ]
    for i, (year, ctype, desc, paid, reserve, status) in enumerate(claims_data):
        claims_table.cell(i + 1, 0).text = year
        claims_table.cell(i + 1, 1).text = ctype
        claims_table.cell(i + 1, 2).text = desc
        claims_table.cell(i + 1, 3).text = paid
        claims_table.cell(i + 1, 4).text = reserve
        claims_table.cell(i + 1, 5).text = status
        for j in range(6):
            for p in claims_table.cell(i + 1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    doc.add_paragraph()
    add_body(doc, '5-Year D&O/Fiduciary Incurred: $11,046,000 (including $4,166,000 defense costs)', bold=True)
    add_body(doc, 'Note: The open securities class action (2025) has reserve of $4.2M against a $15M limit. '
             'If this claim develops adversely, it could consume 28% of the policy limit. Excess carriers '
             '(Chubb Layer 2) have been put on notice.')

    doc.add_paragraph()

    # ── UW Conditions ──
    add_heading_styled(doc, 'Part IX — Special Underwriting Conditions', level=2)
    conditions = [
        '1. SECURITIES CLAIM REPORTING: Any Securities Claim must be reported within 15 days of service of '
        'process on any Insured Person. Late reporting may result in coverage denial.',

        '2. INDEPENDENT COUNSEL: AIG has the right to approve defense counsel selected by the Insured. '
        'Pre-approved panel firms for securities matters: Sullivan & Cromwell, Skadden Arps, Gibson Dunn, '
        'Latham & Watkins. Maximum hourly rates: Partners: $1,450, Associates: $875, Paralegals: $425.',

        '3. ALLOCATION: Where a Claim involves both covered and uncovered matters, or both Insured Persons '
        'and non-insured parties, AIG and the Insured shall use best efforts to agree on a fair and proper '
        'allocation. If unable to agree, the matter shall be submitted to binding arbitration per AAA rules.',

        '4. HAMMER CLAUSE (Modified): If AIG recommends settlement of a Claim and the Insured refuses, '
        'AIG\'s liability shall not exceed: (a) the recommended settlement amount, plus (b) 70% of any '
        'additional Loss incurred above the settlement amount, plus (c) defense costs incurred prior to '
        'the settlement recommendation. The 30/70 split applies (70% AIG, 30% Insured).',

        '5. CHANGE OF CONTROL: If during the Policy Period there is a Change of Control (acquisition of >50% '
        'of voting securities, merger, or asset sale), coverage under this policy shall: (a) continue for '
        'Wrongful Acts occurring prior to the Change of Control date; (b) terminate for Wrongful Acts '
        'occurring after the Change of Control date; and (c) an automatic 6-year Discovery Period shall '
        'apply at no additional premium for Side A only.',

        '6. PENDING MATTERS — SCHEDULE A: The following matters are known to AIG and are excluded from '
        'coverage (or subject to sublimits as noted): (a) Garcia v. Pinnacle — 401(k) fee suit (Fiduciary '
        'sublimit: $2M); (b) Midwest Pension Trust v. Pinnacle — securities class action (currently covered, '
        'reserves established); (c) TCEQ investigation of subsidiary medical waste disposal practices '
        '(not yet a formal Claim — monitoring).',
    ]
    for c in conditions:
        add_body(doc, c)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'This policy is issued on a Claims-Made and Reported basis. Only Claims first made against an '
        'Insured during the Policy Period (or applicable Extended Reporting Period) AND reported to AIG '
        'during the Policy Period (or within 60 days after expiration) are covered. The Retroactive Date '
        'limits coverage to Wrongful Acts occurring on or after January 1, 2018. No coverage is provided '
        'for any Wrongful Act occurring prior to this date.'
    )
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED_RGB
    run.italic = True

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# MAIN
# ============================================================
if __name__ == '__main__':
    print("Generating complex test documents...")
    print()
    generate_workers_comp_retro()
    generate_excess_umbrella()
    generate_cat_loss_report()
    generate_do_fiduciary()
    print()
    print("Done — all complex test documents generated.")
