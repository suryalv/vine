#!/usr/bin/env python3
"""Generate realistic underwriting guideline DOCX documents.

Creates 4 DOCX guideline documents covering major insurance lines.
Each document contains numbered rules with specific, testable thresholds
presented in styled tables — modeled on real carrier guidelines.

Usage:
    python generate_guidelines.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

AIG_BLUE = RGBColor(0x00, 0x20, 0x5B)
BODY_TEXT = RGBColor(0x33, 0x41, 0x55)
MUTED = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def style_header_row(row, bg_color='00205B'):
    """Style a table header row with dark background and white text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = AIG_BLUE
    return h


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


def add_rule_table(doc, headers, rows):
    """Add a styled rule table with header row and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.autofit = True
    for j, h in enumerate(headers):
        table.cell(0, j).text = h
    style_header_row(table.rows[0])
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            table.cell(i + 1, j).text = val
            for p in table.cell(i + 1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
    doc.add_paragraph()
    return table


def add_title_page(doc, title, subtitle_text):
    """Add a styled title and subtitle to the document."""
    title_h = doc.add_heading(title, level=0)
    for run in title_h.runs:
        run.font.color.rgb = AIG_BLUE
    sub = doc.add_paragraph()
    run = sub.add_run(subtitle_text)
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED
    doc.add_paragraph()


# ============================================================
# 1. COMMERCIAL PROPERTY UNDERWRITING GUIDELINES
# ============================================================
def generate_property_guidelines():
    path = os.path.join(OUT_DIR, 'Property_UW_Guidelines.docx')
    doc = Document()

    add_title_page(doc,
                   'Commercial Property Underwriting Guidelines',
                   'AIG North America  |  Effective: January 1, 2026  |  Version 4.2\n'
                   'Classification: INTERNAL USE ONLY  |  Approved by: Chief Underwriting Officer')

    add_body(doc,
             'These guidelines establish the minimum underwriting standards for all commercial '
             'property risks written on AIG paper. All underwriters must comply with these rules '
             'when evaluating new business and renewal submissions. Deviations require written '
             'approval from the Regional Underwriting Manager or above.',
             bold=True)

    # ── SECTION 1: RISK ACCEPTABILITY ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 1: RISK ACCEPTABILITY CRITERIA', level=1)

    add_body(doc,
             'This section defines the construction types, occupancy classes, and risk '
             'characteristics that fall within AIG appetite. Submissions outside these '
             'parameters must be declined or referred to the Risk Acceptability Committee.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('P-101', 'Acceptable Construction Types',
         'ISO Construction Classes 1-6 (Fire Resistive, Modified Fire Resistive, '
         'Masonry Non-Combustible, Non-Combustible, Joisted Masonry, Frame). '
         'Frame construction (Class 1) limited to TIV < $5,000,000.'),
        ('P-102', 'Minimum ISO Class Requirement',
         'ISO Class 4 (Non-Combustible) or higher for any risk with TIV > $25,000,000. '
         'Class 3 (Masonry Non-Combustible) or higher for TIV > $50,000,000.'),
        ('P-103', 'Acceptable Occupancies',
         'Office, warehouse, light manufacturing, retail, institutional, and habitational '
         '(apartments 4+ units). Industrial occupancies acceptable with prior loss review.'),
        ('P-104', 'Prohibited Risks - Absolute',
         'Decline all submissions for: fireworks manufacturing/storage, cannabis cultivation/'
         'processing/dispensary, explosives manufacturing, nuclear facilities, '
         'munitions storage, oil refining (upstream), petrochemical processing, '
         'tire recycling/storage, and wood-frame assisted living > 3 stories.'),
        ('P-105', 'Prohibited Risks - Conditional',
         'The following require VP-level approval: dry cleaning (perc-based), '
         'waste-to-energy facilities, data centers > 50MW, cold storage > 100,000 sq ft, '
         'and any occupancy with spray-applied foam insulation.'),
        ('P-106', 'Building Age Requirements',
         'Buildings > 40 years old require updated appraisal within last 24 months. '
         'Buildings > 60 years old require engineering inspection report and '
         'documentation of all electrical/plumbing/HVAC upgrades.'),
        ('P-107', 'Minimum Protection Class',
         'ISO Public Protection Class (PPC) must be 7 or better. PPC 8 acceptable '
         'only for TIV < $10,000,000. PPC 9-10 require referral with fire department '
         'response time documentation.'),
    ])

    # ── SECTION 2: MINIMUM COVERAGE REQUIREMENTS ──
    add_heading_styled(doc, 'SECTION 2: MINIMUM COVERAGE REQUIREMENTS', level=1)

    add_body(doc,
             'All AIG commercial property policies must meet the following minimum coverage '
             'standards. These requirements protect both the insured and AIG from '
             'underinsurance and coverage disputes.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('P-201', 'Minimum Limits by TIV Tier',
         'TIV $0-$5M: Minimum limit $1,000,000. TIV $5M-$25M: Minimum limit $5,000,000. '
         'TIV $25M-$100M: Minimum limit equal to 80% of TIV. '
         'TIV > $100M: Minimum limit equal to Probable Maximum Loss (PML) as determined by RMS or AIR.'),
        ('P-202', 'Required Coinsurance Percentage',
         'Minimum 80% coinsurance on all Replacement Cost policies. '
         '90% coinsurance required for Agreed Value endorsement eligibility. '
         '100% coinsurance required on blanket policies covering 3+ locations.'),
        ('P-203', 'Required Endorsements - All Risks',
         'All policies must include: Ordinance or Law Coverage (min $500K or 10% of building limit), '
         'Utility Services - Direct Damage, Equipment Breakdown (if not separately placed), '
         'and Debris Removal (min 25% of loss or $250,000).'),
        ('P-204', 'Required Endorsements - By Occupancy',
         'Manufacturing: Mechanical Breakdown, Off-Premises Power Supply. '
         'Habitational: Water Damage Legal Liability (min $100K per unit). '
         'Data Centers: Electronic Data Processing, Media Restoration. '
         'Cold Storage: Spoilage Coverage, Ammonia Contamination.'),
        ('P-205', 'Business Income Minimum Period',
         'All BI coverage must provide minimum 12-month Actual Loss Sustained (ALS) '
         'or agreed amount equal to 50% of annual gross earnings. Extended period of '
         'indemnity: minimum 90 days beyond restoration.'),
        ('P-206', 'Deductible Standards',
         'TIV $0-$5M: Maximum deductible $25,000. TIV $5M-$25M: Maximum deductible $50,000. '
         'TIV $25M-$100M: Maximum deductible $100,000. TIV > $100M: Deductible negotiable, '
         'minimum $100,000. Wind/hail: Percentage deductible required per Section 4.'),
    ])

    # ── SECTION 3: LOSS RATIO STANDARDS ──
    add_heading_styled(doc, 'SECTION 3: LOSS RATIO STANDARDS', level=1)

    add_body(doc,
             'AIG maintains strict loss ratio discipline across the commercial property book. '
             'These standards apply to both new business evaluation and renewal decisions.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('P-301', 'Maximum 5-Year Loss Ratio by Class',
         'Office/Institutional: Decline if 5-year combined ratio > 55%. '
         'Retail/Habitational: Decline if 5-year combined ratio > 60%. '
         'Warehouse/Distribution: Decline if 5-year combined ratio > 65%. '
         'Light Manufacturing: Decline if 5-year combined ratio > 70%. '
         'Heavy Manufacturing: Decline if 5-year combined ratio > 75%.'),
        ('P-302', 'Large Loss Threshold',
         'Any single loss exceeding $500,000 or 25% of annual premium (whichever is less) '
         'triggers mandatory Senior UW review. Two or more large losses in any 3-year period '
         'require Risk Acceptability Committee approval for renewal.'),
        ('P-303', 'Claim Frequency Limits',
         'Decline or non-renew if: >3 property claims in any rolling 12-month period, or '
         '>5 property claims in any rolling 36-month period. Water damage claims: '
         '>2 in any 12-month period triggers mandatory inspection requirement.'),
        ('P-304', 'Incurred-But-Not-Reported (IBNR) Adjustment',
         'When evaluating accounts with open claims, apply IBNR factor: 1.15x for claims '
         '< 12 months old, 1.25x for claims 12-24 months old, 1.10x for claims > 24 months old. '
         'These factors are applied to total incurred for loss ratio calculation.'),
        ('P-305', 'Loss Ratio Trend Analysis',
         'Decline if 3 consecutive policy periods show deteriorating loss ratio AND '
         'the most recent period loss ratio exceeds the class maximum in P-301. '
         'Exception: single catastrophe loss excluded from trend analysis if separately reported.'),
        ('P-306', 'New Business Without Loss History',
         'New ventures (< 3 years operating history) require: minimum 25% rate surcharge, '
         '$50,000 minimum deductible, and 1-year policy term only. '
         'Risk must be classified as "developmental" in the underwriting system.'),
    ])

    # ── SECTION 4: CATASTROPHE EXPOSURE ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 4: CATASTROPHE EXPOSURE', level=1)

    add_body(doc,
             'AIG actively manages catastrophe accumulation across all CAT zones. '
             'These rules establish maximum exposure limits, required deductibles, '
             'and geographic restrictions for wind, flood, and earthquake perils.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('P-401', 'Maximum TIV per CAT Zone - Hurricane',
         'Zone 1 (Coastal TX, LA, FL Panhandle): Max $50M per location, $200M per county. '
         'Zone 2 (Inland Gulf, SE Atlantic coast): Max $75M per location, $300M per county. '
         'Zone 3 (Mid-Atlantic, inland SE): Max $100M per location, $500M per county. '
         'Aggregate limits managed by CAT Management team quarterly.'),
        ('P-402', 'Wind/Hail Deductible Minimums',
         'Tier 1 (Coastal 0-10 mi): 5% of TIV at time of loss, minimum $500,000. '
         'Tier 2 (Near-Coast 10-50 mi): 3% of TIV at time of loss, minimum $250,000. '
         'Tier 3 (Inland 50-100 mi): 2% of TIV at time of loss, minimum $100,000. '
         'Tier 4 (Interior >100 mi): Flat deductible per policy declarations, minimum $50,000.'),
        ('P-403', 'Flood Zone Restrictions',
         'Zone V (Coastal High Hazard): DECLINE. No exceptions. '
         'Zone VE: DECLINE. No exceptions. '
         'Zone AE (100-year floodplain): Maximum $2,500,000 sublimit; require NFIP as primary. '
         'Zone A (100-year, no BFE): Maximum $1,000,000 sublimit; require NFIP and elevation cert. '
         'Zone X Shaded (500-year): Standard terms, no sublimit required. '
         'Zone X Unshaded (Minimal): Standard terms.'),
        ('P-404', 'Earthquake Requirements',
         'CA (fault distance < 5 mi): 15% deductible of TIV, min $250,000. Max TIV $50M. '
         'CA (fault distance 5-25 mi): 10% deductible of TIV, min $100,000. '
         'Pacific NW (Cascadia Zone): 10% deductible of TIV, min $150,000. '
         'New Madrid Seismic Zone: 5% deductible of TIV, min $75,000. '
         'All other zones: 2% deductible of TIV, min $25,000.'),
        ('P-405', 'CAT Aggregate Monitoring',
         'Underwriters must check CAT accumulation dashboard before binding any risk '
         'with TIV > $10,000,000. If county-level accumulation exceeds 80% of maximum, '
         'referral to CAT Management is required. If at 100%, no new risks in that zone.'),
        ('P-406', 'Named Storm Moratorium',
         'Binding authority suspended for all new business and endorsements increasing limits '
         'when NHC issues a Tropical Storm or Hurricane Watch for any county within 200 miles '
         'of the insured location. Moratorium continues until 72 hours after Watch is lifted.'),
    ])

    # ── SECTION 5: PROTECTIVE SAFEGUARDS ──
    add_heading_styled(doc, 'SECTION 5: PROTECTIVE SAFEGUARDS', level=1)

    add_body(doc,
             'Adequate protective safeguards are a condition of coverage. Failure to maintain '
             'required safeguards may void coverage at time of loss per the Protective Safeguards '
             'endorsement (CP 04 11).')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('P-501', 'Sprinkler Requirements by Occupancy',
         'Manufacturing (all types): Full automatic sprinkler system, NFPA 13 compliant. '
         'Warehouse >20,000 sq ft: Full ESFR or in-rack sprinkler per NFPA 13. '
         'Habitational > 3 stories: Full sprinkler per NFPA 13R minimum. '
         'Office > 5 stories: Full sprinkler per NFPA 13. '
         'All other occupancies > $10M TIV: Full sprinkler required.'),
        ('P-502', 'Fire Alarm Requirements',
         'All risks with TIV > $5,000,000: Central station monitored fire alarm system required. '
         'Manufacturing with combustible processes: Addressable fire alarm with smoke/heat detection '
         'in all production areas. Annual inspection certificate must be on file.'),
        ('P-503', 'Roof Maintenance Standards',
         'All risks: Roof inspection report within last 24 months. Flat/built-up roofs > 15 years old: '
         'semi-annual inspection required. Replace if > 20 years old without major renovation. '
         'Metal roofs must be rated for local design wind speed per ASCE 7. '
         'Standing seam preferred over exposed fastener in Tier 1-2 wind zones.'),
        ('P-504', 'Hot Work Procedures',
         'All manufacturing and construction occupancies must have written hot work permit program '
         'compliant with NFPA 51B. Requirements: designated fire watch for minimum 60 minutes '
         'after work completion, pre-work area inspection within 35-foot radius, '
         'and annual training documentation for all employees performing hot work.'),
        ('P-505', 'Electrical Maintenance',
         'All risks with TIV > $10,000,000: Infrared thermographic scan of main electrical panels '
         'within last 24 months. Buildings > 30 years old: require electrical system upgrade '
         'documentation or engineer\'s certification that wiring meets current NEC standards.'),
        ('P-506', 'Water Damage Prevention',
         'Habitational risks > 50 units: Require automatic water shutoff system (leak detection) '
         'in all units. All risks: Domestic water heaters > 12 years old must be replaced. '
         'No flex supply lines older than 5 years. Quarterly inspection of all HVAC condensate '
         'drain lines in multi-story buildings.'),
    ])

    # ── SECTION 6: VALUATION STANDARDS ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 6: VALUATION STANDARDS', level=1)

    add_body(doc,
             'Accurate property valuation is essential to proper risk assessment and adequate '
             'premium development. These standards ensure consistency in valuation methodology '
             'across the AIG property book.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('P-601', 'Replacement Cost vs. Actual Cash Value',
         'Default valuation: Replacement Cost (RC) for all buildings < 30 years old. '
         'ACV permitted only for: buildings > 30 years old awaiting renovation, '
         'buildings scheduled for demolition within 36 months, or at insured\'s written request '
         'with acknowledged coinsurance penalty risk. Functional Replacement Cost: available for '
         'historic structures and religious institutions.'),
        ('P-602', 'Agreed Value Conditions',
         'Agreed Value endorsement available only when: (1) SOV has been updated within 12 months, '
         '(2) independent appraisal by ASA- or AI-designated appraiser is on file for '
         'buildings > $10M value, (3) 90% coinsurance clause is in effect, and '
         '(4) building reconstruction cost estimate uses Marshall & Swift or equivalent methodology.'),
        ('P-603', 'Statement of Values (SOV) Update Frequency',
         'SOVs must be updated: annually for all accounts with TIV > $25M, '
         'every 18 months for accounts with TIV $10M-$25M, '
         'every 24 months for accounts with TIV < $10M. '
         'SOV must include: replacement cost, ACV, square footage, year built, '
         'construction type, occupancy, and protective safeguards for each location.'),
        ('P-604', 'Business Income Calculation Methodology',
         'BI limits must be calculated using the AIG Business Income Worksheet '
         '(Form AIG-BI-100). Methodology: (Annual Gross Revenue - Non-Continuing Expenses) '
         'x (Estimated Period of Restoration in months / 12) x 1.25 safety factor. '
         'Minimum BI limit: 50% of annual gross earnings for manufacturing, '
         '40% for office/retail, 60% for technology/data operations.'),
        ('P-605', 'Equipment and Machinery Valuation',
         'Specialized equipment > $1M individual value: require itemized schedule with '
         'serial numbers, date of manufacture, and individual RC values. '
         'Depreciation schedule: straight-line over manufacturer\'s expected useful life. '
         'EDP equipment: maximum 5-year useful life for valuation purposes.'),
    ])

    # ── SECTION 7: PRICING GUIDELINES ──
    add_heading_styled(doc, 'SECTION 7: PRICING GUIDELINES', level=1)

    add_body(doc,
             'These pricing guidelines establish minimum rates and rating methodologies '
             'for the AIG commercial property book. All rates are subject to state-specific '
             'filing requirements and regulatory constraints.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('P-701', 'Minimum Rates per $100 of TIV by Construction Class',
         'ISO Class 6 (Fire Resistive): $0.08 per $100. '
         'ISO Class 5 (Modified FR): $0.10 per $100. '
         'ISO Class 4 (Non-Combustible): $0.12 per $100. '
         'ISO Class 3 (Masonry Non-Combustible): $0.15 per $100. '
         'ISO Class 2 (Joisted Masonry): $0.20 per $100. '
         'ISO Class 1 (Frame): $0.28 per $100. '
         'These are base rates before ILF, schedule credits, and CAT loads.'),
        ('P-702', 'Loss-Driven Surcharges',
         'Apply surcharge when 5-year loss ratio exceeds 40%: '
         'Loss ratio 40-50%: +10% surcharge. '
         'Loss ratio 50-60%: +20% surcharge. '
         'Loss ratio 60-70%: +35% surcharge. '
         'Loss ratio > 70%: +50% surcharge or decline per P-301. '
         'Surcharges are cumulative with all other rating modifications.'),
        ('P-703', 'Schedule Rating Limits',
         'Maximum schedule credit: -25% (requires documented justification for each factor). '
         'Maximum schedule debit: +50%. Schedule rating factors: construction (+/-5%), '
         'occupancy (+/-5%), protection (+/-10%), maintenance (+/-5%), '
         'management (+/-5%), loss history (+/-15%). '
         'Net schedule modification must be within -25% to +50% range.'),
        ('P-704', 'Catastrophe Load Requirements',
         'Hurricane-exposed risks: Apply RMS/AIR modeled AAL as minimum CAT load. '
         'Earthquake-exposed risks: Apply USGS PGA-based load table. '
         'Minimum CAT load: $500 per $1M of TIV in any CAT zone. '
         'CAT load is added to base rate and is not subject to schedule credits.'),
        ('P-705', 'Minimum Premium',
         'All new commercial property policies: minimum annual premium $2,500. '
         'Accounts with TIV > $10M: minimum annual premium $10,000. '
         'Accounts with TIV > $50M: minimum annual premium $25,000. '
         'Minimum premium is the greater of the calculated premium or these floors.'),
        ('P-706', 'Renewal Pricing Discipline',
         'Maximum renewal rate decrease: -10% without VP approval, -15% with VP approval. '
         'Rate decreases > 15% require Regional UW Manager approval. '
         'No rate decreases permitted if 3-year loss ratio > 50% for the account. '
         'Minimum renewal rate increase when loss ratio > 60%: +10% mandatory.'),
    ])

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'AIG Commercial Property Underwriting Guidelines v4.2 | Effective 01/01/2026 | '
        'Supersedes all prior versions. These guidelines are proprietary and confidential. '
        'Distribution outside AIG is strictly prohibited.')
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# 2. WORKERS' COMPENSATION UNDERWRITING GUIDELINES
# ============================================================
def generate_workers_comp_guidelines():
    path = os.path.join(OUT_DIR, 'Workers_Comp_UW_Guidelines.docx')
    doc = Document()

    add_title_page(doc,
                   'Workers\' Compensation Underwriting Guidelines',
                   'AIG North America  |  Effective: January 1, 2026  |  Version 3.8\n'
                   'Classification: INTERNAL USE ONLY  |  Approved by: Head of Casualty Underwriting')

    add_body(doc,
             'These guidelines establish the minimum underwriting standards for all workers\' '
             'compensation risks written on AIG paper, including guaranteed cost, loss-sensitive, '
             'and retrospectively rated programs. Compliance is mandatory for all underwriters '
             'and deviations require documented approval from the Casualty Line Leader or above.',
             bold=True)

    # ── SECTION 1: RISK SELECTION ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 1: RISK SELECTION', level=1)

    add_body(doc,
             'AIG maintains a defined appetite for workers\' compensation risks. The following '
             'rules govern risk selection, prohibited classes, and minimum eligibility requirements.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('WC-101', 'Acceptable Industry Classes',
         'Acceptable: Professional services (NAICS 54), Finance/Insurance (NAICS 52), '
         'Wholesale trade (NAICS 42), Light manufacturing (NAICS 31-33 excluding heavy metals/chemicals), '
         'Technology (NAICS 51), Healthcare - non-acute (NAICS 621). '
         'Acceptable with restrictions: Construction (NAICS 23) - see WC-102, '
         'Transportation (NAICS 48-49) - max 500 drivers.'),
        ('WC-102', 'Prohibited Class Codes',
         'Auto-decline the following NCCI class codes: '
         '7720 (Underground Mining), 1624 (Quarry Operations), 2702 (Logging/Lumbering), '
         '6251 (Tunneling), 7360 (Carnival/Circus Operations), '
         '3081 (Foundry - Iron), 1463 (Asbestos Contractor), 7403 (Aviation - Commercial Pilot). '
         'Any class code with national loss ratio > 100% over last 5 years: refer to Line Leader.'),
        ('WC-103', 'New Venture Restrictions',
         'Businesses operating < 3 years: maximum $100,000 annual premium. '
         'Businesses operating < 1 year: decline unless (1) principals have 10+ years '
         'industry experience documented, (2) written safety program in place at inception, '
         'and (3) EMR is not available (use 1.00 assumption with +15% surcharge). '
         'No new ventures in prohibited class codes under any circumstances.'),
        ('WC-104', 'Financial Stability Requirements',
         'Standard premium > $100,000: D&B Composite Credit Score >= 50 or equivalent. '
         'Standard premium > $250,000: audited financial statements required. '
         'Standard premium > $500,000: minimum 1.2 current ratio and 0.8 debt-to-equity maximum. '
         'Companies with active bankruptcy proceedings: decline.'),
        ('WC-105', 'Minimum Premium Requirements',
         'Guaranteed cost: minimum annual premium $5,000 or applicable state minimum, '
         'whichever is greater. Large deductible program: minimum $250,000 standard premium. '
         'Retrospectively rated program: minimum $250,000 standard premium (see Section 3).'),
        ('WC-106', 'Multi-Year Policy Restrictions',
         'Maximum policy term: 1 year for new business. '
         'Renewal accounts with 3+ year tenure and loss ratio < 50%: eligible for 3-year rate guarantee. '
         'No multi-year policies for any account with EMR > 1.20 or loss ratio > 65%.'),
    ])

    # ── SECTION 2: EXPERIENCE MODIFICATION ──
    add_heading_styled(doc, 'SECTION 2: EXPERIENCE MODIFICATION', level=1)

    add_body(doc,
             'The Experience Modification Rate (EMR) is the primary indicator of a risk\'s '
             'relative loss performance. AIG uses EMR thresholds tiered by premium size to '
             'account for the credibility of the modification factor.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('WC-201', 'EMR Acceptability by Premium Size',
         'Small accounts (standard premium < $100K): Accept if EMR <= 1.25. '
         'Refer if EMR 1.26-1.40. Decline if EMR > 1.40. '
         'Medium accounts ($100K-$500K): Accept if EMR <= 1.35. '
         'Refer if EMR 1.36-1.60. Decline if EMR > 1.60. '
         'Large accounts (> $500K): Accept if EMR <= 1.50. '
         'Refer if EMR 1.51-1.75. Decline if EMR > 1.75.'),
        ('WC-202', 'EMR Trend Analysis',
         'Obtain 3-year EMR history for all accounts. Decline if EMR has increased '
         'in each of the last 3 consecutive years AND current EMR exceeds the accept '
         'threshold in WC-201. Exception: documented corrective safety measures '
         'implemented within last 12 months may be considered for referral rather than decline.'),
        ('WC-203', 'Auto-Decline EMR Thresholds',
         'Automatically decline any risk with EMR > 2.00 regardless of premium size. '
         'Automatically decline any risk with EMR > 1.50 in prohibited-adjacent class codes '
         '(any class code within the same 3-digit classification group as a prohibited code). '
         'No exceptions without written approval from VP of Casualty.'),
        ('WC-204', 'Interstate EMR Considerations',
         'For multi-state risks, evaluate the Interstate EMR (if available) as primary indicator. '
         'If state-specific EMRs vary by more than 0.30 from the interstate mod, '
         'evaluate each state independently and apply the highest state EMR for '
         'underwriting threshold purposes under WC-201.'),
        ('WC-205', 'EMR Not Available',
         'When EMR is not available (new ventures, insufficient premium volume): '
         'Assume EMR of 1.00 for pricing purposes. Apply +15% new venture surcharge. '
         'Require loss runs from prior carrier(s) for any experience period available. '
         'If no prior WC coverage existed, require written explanation and evidence of '
         'prior compliance with state WC requirements.'),
    ])

    # ── SECTION 3: RETROSPECTIVE RATING ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 3: RETROSPECTIVE RATING', level=1)

    add_body(doc,
             'Retrospectively rated programs are available for qualified risks meeting '
             'premium and financial requirements. These guidelines establish minimum '
             'program parameters and collateral standards.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('WC-301', 'Eligibility Requirements',
         'Minimum standard premium: $250,000 annual. Minimum years in business: 5. '
         'EMR must be <= 1.20. Loss ratio (5-year average) must be <= 60%. '
         'D&B Composite Credit Score >= 60 or audited financials demonstrating adequate liquidity. '
         'Insured must have dedicated risk management staff or retained safety consultant.'),
        ('WC-302', 'Maximum/Minimum Premium Factors',
         'Minimum premium factor: Not less than 0.45 (45% of standard premium). '
         'Maximum premium factor: Standard is 1.50 (150% of standard premium). '
         'For accounts with premium > $1M: max factor negotiable to 2.00 with collateral. '
         'Swing range (max - min) must not exceed 1.20 for accounts < $500K standard premium.'),
        ('WC-303', 'Collateral Requirements',
         'Cash collateral or irrevocable letter of credit required for all retro programs. '
         'Minimum collateral: estimated outstanding losses + 50% of unearned premium factor swing. '
         'Collateral must be posted within 30 days of binding. '
         'Failure to post collateral within 60 days: convert to guaranteed cost at standard rates. '
         'Annual collateral review and adjustment at each retro adjustment date.'),
        ('WC-304', 'Loss Conversion Factor (LCF) Standards',
         'Standard LCF: 1.12 for guaranteed cost equivalent. '
         'ALAE included in loss: LCF = 1.00 (allocated expense within loss limit). '
         'ALAE excluded from loss: LCF = 1.12 standard, 1.08 for accounts > $1M premium. '
         'LCF is applied to limited losses before calculating the retro premium.'),
        ('WC-305', 'Per-Occurrence Loss Limitations',
         'Standard per-occurrence limit: $250,000 for accounts with $250K-$500K premium. '
         '$350,000 for accounts with $500K-$1M premium. '
         '$500,000 for accounts with $1M-$2M premium. '
         'Negotiable above $500,000 for accounts with premium > $2M. '
         'Unlimited loss inclusion requires VP approval and additional collateral.'),
        ('WC-306', 'Retro Adjustment Schedule',
         'First adjustment: 18 months from policy inception. '
         'Subsequent adjustments: annually until closure. '
         'Maximum adjustment period: 6 years from policy expiration. '
         'Final close-out: 7 years from expiration or when all claims closed, whichever is first.'),
    ])

    # ── SECTION 4: SAFETY REQUIREMENTS ──
    add_heading_styled(doc, 'SECTION 4: SAFETY REQUIREMENTS', level=1)

    add_body(doc,
             'AIG requires documented safety programs and performance standards for all '
             'workers\' compensation risks. These requirements protect employees and '
             'demonstrate the insured\'s commitment to loss prevention.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('WC-401', 'Required Safety Programs by Industry',
         'Construction: Written safety program including fall protection, scaffolding, '
         'trenching/excavation, and confined space entry protocols per OSHA 29 CFR 1926. '
         'Manufacturing: Lock-out/tag-out, machine guarding, PPE, and hazard communication programs. '
         'Healthcare: Needlestick prevention, patient handling (lift equipment), and workplace violence programs. '
         'Transportation: Driver qualification, hours of service monitoring, vehicle maintenance, '
         'and MVR review program (annual for all drivers).'),
        ('WC-402', 'OSHA Incident Rate Thresholds',
         'Total Recordable Incident Rate (TRIR): Decline if TRIR > 2.0x the industry average '
         'published by BLS for the applicable NAICS code. Refer if TRIR > 1.5x industry average. '
         'Days Away, Restricted, or Transferred (DART) rate: Decline if DART > 2.0x industry average. '
         'Obtain 3-year OSHA 300A logs for all accounts with premium > $50,000.'),
        ('WC-403', 'Return-to-Work Program Standards',
         'All accounts with premium > $100,000: must have written return-to-work (RTW) program. '
         'RTW program must include: modified duty job descriptions, maximum modified duty period '
         '(recommend 90 days), physician communication protocol, and employee notification procedures. '
         'Accounts without RTW program: apply +5% surcharge.'),
        ('WC-404', 'Safety Inspection Requirements',
         'New business: Pre-bind loss control survey required for all accounts > $150,000 premium. '
         'Renewal: Annual loss control survey for accounts > $250,000 premium. '
         'Critical recommendations from loss control must be completed within 90 days '
         'or coverage may be non-renewed. Open critical recommendations at renewal: +10% surcharge.'),
        ('WC-405', 'Drug and Alcohol Testing',
         'Construction (all class codes): Pre-employment and post-accident drug testing required. '
         'Transportation (all drivers): DOT-compliant random drug and alcohol testing required. '
         'All other industries: Pre-employment drug testing recommended. '
         'Accounts with substance abuse-related claims: mandatory testing program required.'),
    ])

    # ── SECTION 5: CLAIMS MANAGEMENT ──
    add_heading_styled(doc, 'SECTION 5: CLAIMS MANAGEMENT', level=1)

    add_body(doc,
             'Effective claims management is critical to maintaining acceptable loss ratios '
             'and protecting injured workers. These standards apply to underwriting evaluation '
             'of an insured\'s claims performance.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('WC-501', 'Claim Frequency Thresholds',
         'Evaluate claim frequency per 100 FTEs. Decline if frequency > 15 claims per 100 FTEs '
         'for office/professional classes. Decline if frequency > 25 claims per 100 FTEs for '
         'light manufacturing/warehouse. Decline if frequency > 35 claims per 100 FTEs for '
         'heavy manufacturing/construction. Frequency must show stable or declining trend '
         'over 3-year evaluation period.'),
        ('WC-502', 'Serious Injury/Fatality (SIF) Review',
         'Any account with 1 or more fatalities in the last 5 years: mandatory referral to '
         'VP of Casualty. Any account with 2+ permanent total disability claims in 5 years: '
         'decline unless corrective actions documented and verified by AIG Loss Control. '
         'All SIF claims must have OSHA investigation report on file before quoting.'),
        ('WC-503', 'Subrogation Protocols',
         'All claims with potential third-party liability: subrogation investigation required '
         'within 30 days of claim notification. Minimum subrogation recovery target: 8% of '
         'total incurred for construction risks, 5% for manufacturing, 3% for all others. '
         'Subrogation performance is factored into renewal pricing as a credit (up to -5%).'),
        ('WC-504', 'Medical Cost Containment',
         'All accounts must utilize AIG Preferred Provider Organization (PPO) network where available. '
         'Utilization review required for all inpatient stays and surgical procedures. '
         'Pharmacy benefit management program mandatory for all claims with prescription costs '
         'exceeding $2,500. Evaluate insured\'s medical cost per claim vs. industry benchmark.'),
        ('WC-505', 'Litigation Rate Analysis',
         'Maximum acceptable attorney involvement rate: 20% of total claims. '
         'If litigation rate > 30%: require written action plan from insured and apply +8% surcharge. '
         'If litigation rate > 40%: decline or non-renew. '
         'Evaluate jurisdiction-specific litigation trends and defense costs.'),
    ])

    # ── SECTION 6: MULTI-STATE OPERATIONS ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 6: MULTI-STATE OPERATIONS', level=1)

    add_body(doc,
             'Workers\' compensation is state-regulated with significant jurisdictional '
             'variation. These guidelines address multi-state complexities, monopolistic '
             'state requirements, and federal program considerations.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('WC-601', 'Monopolistic State Handling',
         'Ohio, North Dakota, Washington, and Wyoming are monopolistic states where WC coverage '
         'must be obtained from the state fund. AIG cannot write primary WC in these states. '
         'Stop-gap / Employers Liability coverage: Available via endorsement on the master policy. '
         'Stop-gap limit: match EL limits on the master policy, minimum $500,000 each accident.'),
        ('WC-602', 'United States Longshore & Harbor Workers (USL&H)',
         'USL&H coverage required for all employees working on or over navigable waters '
         'per 33 USC 901-950. Minimum USL&H limits: $1,000,000 each accident / $1,000,000 '
         'disease-policy limit / $1,000,000 disease-each employee. '
         'Maritime employers must also carry Jones Act coverage (separate quotation required). '
         'USL&H class code surcharge: +25% above state standard premium rates.'),
        ('WC-603', 'Foreign Voluntary Workers\' Compensation',
         'Available for U.S. employees temporarily working abroad (< 180 days per assignment). '
         'Territory: Worldwide excluding sanctioned countries (see current OFAC SDN list). '
         'Coverage limit: $1,000,000 per accident. Repatriation expense: $250,000 sublimit. '
         'Endemic disease: $100,000 sublimit. War risk and terrorism: excluded. '
         'Employees permanently assigned overseas: separate expat program required.'),
        ('WC-604', 'Federal Programs',
         'Federal Employees\' Compensation Act (FECA): Not available on AIG paper. '
         'Defense Base Act (DBA): Available for government contractors; separate quotation. '
         'Federal Coal Mine Safety Act: Decline all submissions. '
         'Nuclear Workers: Decline all submissions (refer to AIG Nuclear pool).'),
        ('WC-605', 'State-Specific Underwriting Alerts',
         'California: Require quarterly payroll reporting for all accounts > $200K premium. '
         'New York: Apply NY Construction Employment Act requirements for GC/sub relationships. '
         'Florida: Construction class codes require proof of subcontractor WC certificates. '
         'Illinois: Medical fee schedule changes effective 07/01/2026 - adjust reserves accordingly. '
         'Pennsylvania: Uninsured employer fund surcharge of 3% applies.'),
    ])

    # ── SECTION 7: PRICING STANDARDS ──
    add_heading_styled(doc, 'SECTION 7: PRICING STANDARDS', level=1)

    add_body(doc,
             'Workers\' compensation pricing must comply with state-filed rates and rating '
             'plans. These standards ensure consistent application of available rating modifications.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('WC-701', 'Minimum Premium by State',
         'Apply the greater of: (1) state-filed minimum premium for each applicable class code, '
         'or (2) AIG minimum premium of $5,000 per state. '
         'Multi-state policies: minimum premium applies to each state individually. '
         'If total policy premium < $10,000 across all states: decline (below AIG minimum threshold).'),
        ('WC-702', 'Schedule Rating Credit Limits',
         'Maximum schedule credit: per state-filed schedule rating plan (typically -25% to -40%). '
         'AIG internal maximum: -25% without Line Leader approval, -35% with Line Leader approval. '
         'Credits exceeding state maximum: prohibited under all circumstances. '
         'All schedule credits must be documented with specific loss control justification '
         'for each credit category (premises, classification, management, safety, etc.).'),
        ('WC-703', 'Premium Discount Tables',
         'Apply NCCI or state-specific premium discount for eligible accounts: '
         '$5,000-$10,000 premium: 5.5% discount. $10,001-$100,000: 9.0% discount. '
         '$100,001-$500,000: 11.5% discount. $500,001-$1,000,000: 13.0% discount. '
         '$1,000,001+: 14.0% discount. '
         'Premium discount is applied after experience rating and schedule rating modifications.'),
        ('WC-704', 'Expense Constant',
         'Apply state-filed expense constant to all policies. Standard NCCI expense constant: '
         '$250 per policy. State-specific expense constants vary; use applicable state filing. '
         'Expense constant is not subject to experience modification or schedule rating.'),
        ('WC-705', 'Retrospective Rating Premium Tax Multiplier',
         'Apply state-specific premium tax multiplier to all retro premium computations. '
         'Standard multiplier range: 1.03 to 1.07 depending on state. '
         'Include assessment factors, second-injury fund surcharges, and other state surcharges '
         'in the tax multiplier. Update annually per state filing.'),
        ('WC-706', 'Renewal Rate Cap',
         'Guaranteed cost renewals: maximum +15% rate increase without documented loss justification. '
         'Rate increases > 25% require 60-day advance notice and VP approval. '
         'No rate decreases permitted if 3-year loss ratio > 70% for the account. '
         'Rate decreases > 10% require written justification filed in underwriting workbook.'),
    ])

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'AIG Workers\' Compensation Underwriting Guidelines v3.8 | Effective 01/01/2026 | '
        'Supersedes all prior versions. These guidelines are proprietary and confidential. '
        'Distribution outside AIG is strictly prohibited.')
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# 3. EXCESS / UMBRELLA LIABILITY UNDERWRITING GUIDELINES
# ============================================================
def generate_excess_umbrella_guidelines():
    path = os.path.join(OUT_DIR, 'Excess_Umbrella_UW_Guidelines.docx')
    doc = Document()

    add_title_page(doc,
                   'Excess / Umbrella Liability Underwriting Guidelines',
                   'AIG North America  |  Effective: January 1, 2026  |  Version 5.1\n'
                   'Classification: INTERNAL USE ONLY  |  Approved by: Head of Excess Casualty')

    add_body(doc,
             'These guidelines govern all excess and umbrella liability placements, including '
             'lead umbrella, excess follow-form, and quota share participations. All underwriters '
             'must verify underlying coverage structure and carrier quality before binding any '
             'excess or umbrella layer. Deviations from these guidelines require Excess Casualty '
             'Line Leader approval.',
             bold=True)

    # ── SECTION 1: UNDERLYING REQUIREMENTS ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 1: UNDERLYING REQUIREMENTS', level=1)

    add_body(doc,
             'AIG excess/umbrella policies sit above a defined schedule of underlying coverage. '
             'The adequacy and quality of the underlying program directly impacts the excess '
             'layer\'s exposure. These minimum standards are non-negotiable.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('EU-101', 'Minimum Underlying Limits - Commercial General Liability',
         'CGL: $1,000,000 per occurrence / $2,000,000 general aggregate / $2,000,000 '
         'products-completed operations aggregate / $1,000,000 personal & advertising injury. '
         'Must include contractual liability, independent contractors, and '
         'products-completed operations coverage. XCU exclusions: not acceptable.'),
        ('EU-102', 'Minimum Underlying Limits - Commercial Auto',
         'Business Auto: $1,000,000 combined single limit (CSL) for bodily injury and '
         'property damage. Must include hired and non-owned auto coverage. '
         'For fleets > 100 vehicles or any vehicle > 26,001 lbs GVW: minimum $2,000,000 CSL. '
         'MCS-90 endorsement required for all motor carriers.'),
        ('EU-103', 'Minimum Underlying Limits - Employers Liability',
         'Workers\' Compensation Employers Liability: $1,000,000 each accident / '
         '$1,000,000 disease-policy limit / $1,000,000 disease-each employee. '
         'Stop-gap coverage required for employees in monopolistic states. '
         'USL&H coverage required for applicable maritime exposures.'),
        ('EU-104', 'Minimum Underlying Limits - Other Lines',
         'If scheduled as underlying: Professional Liability: $1,000,000 per claim / $2,000,000 aggregate. '
         'Liquor Liability: $1,000,000 per occurrence (if applicable). '
         'Employee Benefits Liability: $1,000,000 per occurrence. '
         'Watercraft Liability: $1,000,000 per occurrence (if applicable).'),
        ('EU-105', 'Underlying Carrier Quality Standards',
         'All underlying carriers must maintain A.M. Best rating of A- (Excellent) VII '
         'or higher. If underlying carrier is downgraded below A- VII during the policy period, '
         'insured must replace the carrier within 90 days or AIG may issue coverage gap notice. '
         'Self-insured retentions > $250,000: require actuarial certification of adequacy.'),
        ('EU-106', 'Defense Cost Provisions',
         'AIG lead umbrella: Defense costs are outside the limit (supplementary payments). '
         'AIG excess follow-form: Defense costs follow the underlying treatment. '
         'If underlying provides defense within limits: AIG excess limit will erode '
         'proportionally and insured must be advised in writing. '
         'Duty to defend: AIG has no duty to defend in excess layers above $5M attachment.'),
        ('EU-107', 'Underlying Coverage Verification',
         'Copies of all scheduled underlying policies or binders must be on file before binding. '
         'Evidence of insurance (certificates) are not acceptable substitutes for policy copies. '
         'Verify all underlying policies have concurrent inception/expiration dates. '
         'Any gap in underlying coverage > 1 day: require written explanation and bridge coverage.'),
    ])

    # ── SECTION 2: ATTACHMENT POINTS ──
    add_heading_styled(doc, 'SECTION 2: ATTACHMENT POINTS', level=1)

    add_body(doc,
             'Attachment point selection directly impacts loss frequency and severity for the '
             'excess layer. These guidelines establish minimum attachment requirements by '
             'risk class and program structure.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('EU-201', 'Minimum Attachment by Industry',
         'Low hazard (office, retail, professional services): $1,000,000 minimum attachment. '
         'Moderate hazard (light manufacturing, warehouse, restaurants): $2,000,000 minimum attachment. '
         'High hazard (heavy construction, transportation, chemical manufacturing): '
         '$5,000,000 minimum attachment. '
         'Healthcare (hospitals, long-term care): $5,000,000 minimum attachment.'),
        ('EU-202', 'Maximum Attachment for AIG Lead',
         'AIG will lead (provide first excess layer) only where attachment point does not exceed '
         '$25,000,000 for Fortune 500 accounts, $10,000,000 for middle market accounts, '
         'and $5,000,000 for small commercial accounts. '
         'Attachments above these thresholds: AIG participates as follow-form only.'),
        ('EU-203', 'Follow-Form vs. Independent Coverage',
         'Follow-form (excess of underlying): Available when underlying is an approved form '
         '(ISO CGL, AIG Primary CGL, or equivalent). Coverage mirrors underlying with stated exceptions. '
         'Independent coverage (standalone umbrella form): Required when underlying includes '
         'non-standard exclusions or when AIG cannot verify underlying coverage adequacy. '
         'Independent form: AIG Umbrella Form UMB-2026 applies.'),
        ('EU-204', 'Gap in Underlying Coverage',
         'If any scheduled underlying coverage has a coverage gap (exclusion not present in '
         'AIG umbrella form): AIG umbrella drops down to provide primary coverage for the gap '
         'subject to the underlying per-occurrence limit as a self-insured retention. '
         'Drop-down not available for: pollution, employment practices, professional liability, '
         'or cyber exposures unless specifically endorsed.'),
    ])

    # ── SECTION 3: EXCLUDED RISKS ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 3: EXCLUDED RISKS', level=1)

    add_body(doc,
             'Certain risks are excluded from AIG excess/umbrella coverage either absolutely '
             '(no exceptions) or conditionally (available with restrictions). Underwriters must '
             'verify the applicability of each exclusion to every submission.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('EU-301', 'Absolute Exclusions',
         'The following are excluded under all circumstances, no exceptions or buybacks available: '
         'PFAS / PFOS / PFOA (per- and polyfluoroalkyl substances) contamination. '
         'Cannabis / marijuana operations (cultivation, processing, distribution, dispensary). '
         'Nuclear energy liability (refer to AIG Nuclear pool). '
         'War, invasion, armed conflict, or military action. '
         'Government-imposed sanctions (OFAC compliance). '
         'Asbestos (known exposure at inception). '
         'Silica and silica-related disease. '
         'Lead paint in residential structures built before 1978.'),
        ('EU-302', 'Conditional Exclusions - With Sublimit',
         'The following may be covered with specific sublimits and additional premium: '
         'AI / Autonomous Systems Liability: Maximum $2,000,000 sublimit within the layer, '
         'requires AI risk questionnaire and technology review. '
         'Unmanned Aircraft Systems (Drones): Maximum $1,000,000 sublimit, requires FAA Part 107 compliance. '
         'Cyber Liability (bodily injury/property damage only): Maximum $1,000,000 sublimit.'),
        ('EU-303', 'Employment Practices Gap Treatment',
         'Employment practices liability (EPL) is excluded under the standard umbrella form. '
         'If underlying EPL coverage exists with limits less than umbrella attachment: '
         'gap is NOT covered by umbrella unless EPL Buyback endorsement is issued. '
         'EPL Buyback endorsement: available for additional premium, maximum $5,000,000 sublimit, '
         'requires dedicated EPL underlying policy with minimum $1M limits.'),
        ('EU-304', 'Professional Liability Treatment',
         'Professional liability / errors & omissions is excluded under standard umbrella form. '
         'Excess professional liability: available as separate placement on AIG XPL form. '
         'Miscellaneous professional liability buyback: available for non-professional service '
         'businesses with incidental professional exposure (maximum $2,000,000 sublimit).'),
        ('EU-305', 'Pollution Exclusion',
         'Total pollution exclusion applies to all excess/umbrella policies. '
         'Hostile fire exception: pollution resulting from a fire that is hostile (unintended) '
         'is covered to the extent the fire itself would be covered. '
         'Sudden and accidental pollution buyback: not available on excess/umbrella layers. '
         'Separate environmental liability placement required for any pollution exposure.'),
    ])

    # ── SECTION 4: TOWER STRUCTURE ──
    add_heading_styled(doc, 'SECTION 4: TOWER STRUCTURE', level=1)

    add_body(doc,
             'AIG manages its exposure across multi-layer excess programs through participation '
             'limits, co-insurance requirements, and drop-down provisions. These rules ensure '
             'balanced tower construction.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('EU-401', 'Maximum Participation per Layer',
         'Lead umbrella ($1M x $1M): AIG maximum 100% participation. '
         'First excess ($5M x $1M to $5M x $5M): AIG maximum 100% participation. '
         'Second excess ($10M x $10M): AIG maximum 50% participation, require co-carrier(s). '
         'Higher excess layers (> $25M attachment): AIG maximum 25% participation. '
         'Total AIG participation across all layers: Maximum $50M for any single insured.'),
        ('EU-402', 'Co-Insurance Requirements',
         'For any layer where AIG participates < 100%: require minimum 2 co-carriers. '
         'All co-carriers must maintain A.M. Best rating A- VII or higher. '
         'Lead carrier on any layer must take minimum 25% participation. '
         'No carrier may participate on a layer for less than 10% share.'),
        ('EU-403', 'Drop-Down Provisions',
         'AIG excess layers will drop down in the event of: exhaustion of underlying aggregate '
         '(upon proof of payment by underlying carrier), insolvency of an underlying carrier '
         '(subject to AIG\'s right to rescind within 60 days of insolvency notice), '
         'or a coverage gap where AIG form is broader than underlying. '
         'Drop-down does not apply to: intentional acts, criminal conduct, or punitive damages.'),
        ('EU-404', 'Punitive Damage Territory Restrictions',
         'Punitive damages are EXCLUDED in the following states where they are deemed uninsurable '
         'as against public policy: Colorado, Connecticut, Illinois, Indiana, Kansas, Louisiana, '
         'Michigan, Minnesota, Nebraska, New York, Oregon, Pennsylvania, Utah, and Virginia. '
         'Punitive damages covered where insurable by law. '
         'Most favorable jurisdiction clause: NOT available on AIG excess forms.'),
        ('EU-405', 'Erosion and Reinstatement',
         'General aggregate: 2x per occurrence limit standard. '
         'Products-completed operations aggregate: 2x per occurrence limit standard. '
         'No automatic reinstatement of aggregate after exhaustion. '
         'Reinstatement available by endorsement at 100% additional premium pro-rata for '
         'remaining term. Maximum one reinstatement per policy period.'),
    ])

    # ── SECTION 5: LOSS HISTORY ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 5: LOSS HISTORY', level=1)

    add_body(doc,
             'Loss history analysis for excess/umbrella risks focuses on severity, defense cost '
             'trends, and the potential for claims to penetrate the proposed excess layer.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('EU-501', 'Maximum Acceptable Loss Ratio',
         'Maximum 5-year excess loss ratio (losses penetrating proposed layer / premium): 35%. '
         'Decline if excess loss ratio > 50% in any single year of the last 5 years. '
         'For lead umbrella: maximum total loss ratio (ground-up) of 60% over 5 years. '
         'Loss ratio calculation must include ALAE as reported.'),
        ('EU-502', 'Large Loss Review Triggers',
         'Mandatory review by Excess Casualty VP for any of the following: '
         'Single occurrence loss > $1,000,000 ground-up. '
         'Single occurrence loss that has penetrated or is projected to penetrate any excess layer. '
         'Defense costs on a single claim exceeding $500,000. '
         'Any claim involving fatality or permanent total disability.'),
        ('EU-503', 'Defense Cost Analysis',
         'Evaluate defense costs as a percentage of total incurred: maximum acceptable ratio of '
         'defense costs to indemnity is 40% for general liability, 50% for auto liability, '
         'and 60% for products liability. Excessive defense cost ratios indicate inefficient '
         'claim handling and trigger recommendation to require panel counsel approval.'),
        ('EU-504', 'Loss Projection Requirements',
         'For all accounts with > $1M premium: require actuarial loss projection from '
         'approved actuarial firm. Projection must include: expected ultimate losses at '
         'proposed attachment, frequency and severity distributions, and '
         'probability of layer attachment within the policy period. '
         'Minimum confidence level for pricing: 80th percentile of projected loss distribution.'),
    ])

    # ── SECTION 6: INTERNATIONAL ──
    add_heading_styled(doc, 'SECTION 6: INTERNATIONAL', level=1)

    add_body(doc,
             'AIG excess/umbrella policies may provide worldwide coverage territory subject to '
             'the restrictions and requirements in this section.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('EU-601', 'Coverage Territory Restrictions',
         'Standard coverage territory: Worldwide, but suits must be brought in the United States, '
         'its territories, possessions, or Canada. International suits covered only via '
         'Worldwide Coverage endorsement (additional premium required). '
         'EU/EEA jurisdiction claims: covered with Worldwide endorsement. '
         'All other jurisdictions: require specific country-by-country approval.'),
        ('EU-602', 'Sanctioned Countries',
         'No coverage available for operations, activities, or claims arising in or from: '
         'Cuba, Iran, North Korea, Syria, and the Crimea/Donetsk/Luhansk regions of Ukraine. '
         'Additional countries per current OFAC Sanctions Programs list. '
         'Underwriters must check OFAC SDN list before binding any account with '
         'international operations. Automatic exclusion endorsement applies.'),
        ('EU-603', 'Difference in Conditions / Difference in Limits (DIC/DIL)',
         'For insureds with operations in admitted-paper jurisdictions: AIG master policy '
         'provides DIC/DIL coverage to fill gaps between local admitted policies and the '
         'master program. DIC: covers perils included in master but excluded from local policy. '
         'DIL: provides excess limits above local policy limits up to master policy limit. '
         'DIC/DIL is not available where prohibited by local law or regulation.'),
        ('EU-604', 'Local Admitted Paper Requirements',
         'Insureds with permanent operations, employees, or assets in foreign jurisdictions '
         'must maintain local admitted liability policies where required by law. '
         'AIG Global Network: coordinate through AIG Multinational team for local policy issuance. '
         'Countries requiring admitted paper: all EU member states, Brazil, Mexico, China, '
         'India, Japan, Australia, and any country where non-admitted insurance is prohibited. '
         'Penalty for non-compliance is on the insured; AIG will not indemnify regulatory fines.'),
        ('EU-605', 'Foreign Jurisdiction Claims Handling',
         'Claims arising in foreign jurisdictions must be reported to AIG International Claims '
         'within 15 days of insured\'s knowledge. Local counsel must be engaged within 30 days. '
         'AIG retains right to appoint and direct defense counsel in all jurisdictions. '
         'Settlement authority follows the master policy terms regardless of local jurisdiction.'),
    ])

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'AIG Excess / Umbrella Liability Underwriting Guidelines v5.1 | Effective 01/01/2026 | '
        'Supersedes all prior versions. These guidelines are proprietary and confidential. '
        'Distribution outside AIG is strictly prohibited.')
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# 4. D&O AND FIDUCIARY LIABILITY UNDERWRITING GUIDELINES
# ============================================================
def generate_do_fiduciary_guidelines():
    path = os.path.join(OUT_DIR, 'DO_Fiduciary_UW_Guidelines.docx')
    doc = Document()

    add_title_page(doc,
                   'Directors & Officers and Fiduciary Liability\nUnderwriting Guidelines',
                   'AIG Financial Lines  |  Effective: January 1, 2026  |  Version 6.0\n'
                   'Classification: INTERNAL USE ONLY  |  Approved by: Head of Financial Lines')

    add_body(doc,
             'These guidelines apply to all Directors & Officers (D&O), Employment Practices '
             'Liability (EPLI), and Fiduciary Liability policies written by AIG Financial Lines. '
             'The guidelines cover public company D&O, private/not-for-profit D&O, Side A DIC, '
             'and fiduciary liability placements. All submissions must be evaluated against '
             'these standards before quotation.',
             bold=True)

    # ── SECTION 1: RISK CLASSIFICATION ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 1: RISK CLASSIFICATION', level=1)

    add_body(doc,
             'D&O risk characteristics vary significantly based on company type, size, '
             'industry, and financial health. These classification standards ensure proper '
             'risk segmentation and pricing adequacy.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('DO-101', 'Public vs. Private Company Criteria',
         'Public company D&O: Any entity with securities registered under the Securities '
         'Exchange Act of 1934 (SEC reporting), including: NYSE/NASDAQ-listed companies, '
         'OTC-traded companies, SPACs (pre- and post-de-SPAC), ADR issuers with U.S. listings. '
         'Private company D&O: All other entities including privately held corporations, '
         'LLCs, partnerships, not-for-profit organizations, and privately held subsidiaries '
         'of public companies.'),
        ('DO-102', 'Market Capitalization Tiers and Limits',
         'Micro-cap (< $300M market cap): Maximum $10M D&O limit per policy. '
         'Small-cap ($300M-$2B): Maximum $15M D&O limit per policy. '
         'Mid-cap ($2B-$10B): Maximum $25M D&O limit per policy. '
         'Large-cap ($10B-$100B): Maximum $50M D&O limit per policy. '
         'Mega-cap (> $100B): Maximum $25M AIG participation; tower structure required.'),
        ('DO-103', 'Revenue-Based Retention Minimums',
         'Revenue < $100M: Minimum retention $100,000 (Securities Entity, Side B/C). '
         'Revenue $100M-$500M: Minimum retention $250,000. '
         'Revenue $500M-$1B: Minimum retention $500,000. '
         'Revenue $1B-$5B: Minimum retention $1,000,000. '
         'Revenue $5B-$10B: Minimum retention $2,500,000. '
         'Revenue > $10B: Minimum retention $5,000,000. '
         'Side A (individual director/officer only): No retention.'),
        ('DO-104', 'Industry Classification Risk Tiers',
         'Tier 1 (Favorable): Technology (SaaS/enterprise software), professional services, '
         'consumer staples. Tier 2 (Standard): Manufacturing, retail, transportation, '
         'real estate (non-REIT). Tier 3 (Elevated): Biotech/pharma (clinical stage), '
         'financial services, cryptocurrency/blockchain, SPACs. '
         'Tier 4 (High Risk): Cannabis-related SPACs, Chinese reverse mergers (VIE structures), '
         'de-SPAC within 24 months, IPO within 12 months. '
         'Tier 4 risks: refer to Financial Lines VP, minimum +50% rate surcharge.'),
        ('DO-105', 'Not-for-Profit Classification',
         'Standard NFP: Charitable organizations, trade associations, private foundations '
         'with annual budget < $100M. Maximum D&O limit: $10M. '
         'Large NFP: Healthcare systems, universities, major charities with budget > $100M. '
         'Underwrite as comparable private company with fiduciary overlay. '
         'Government-affiliated entities: decline (sovereign immunity considerations).'),
    ])

    # ── SECTION 2: COVERAGE STRUCTURE ──
    add_heading_styled(doc, 'SECTION 2: COVERAGE STRUCTURE', level=1)

    add_body(doc,
             'D&O coverage is structured in multiple insuring agreements (Sides A, B, and C). '
             'Each side responds to different types of loss and has distinct underwriting '
             'considerations.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('DO-201', 'Side A Coverage Parameters',
         'Side A (Non-Indemnifiable Loss): Covers individual directors and officers when the '
         'company cannot or will not indemnify. No retention applies. '
         'Required on all public company D&O placements. '
         'Side A limit: may be equal to or greater than Side B/C limit. '
         'Side A priority: always primary to and non-contributory with Side B/C.'),
        ('DO-202', 'Side B/C Coverage Parameters',
         'Side B (Corporate Reimbursement): Covers the entity when it indemnifies D&Os. '
         'Side C (Entity Securities): Covers the entity for securities claims (public companies only). '
         'Side C limit: shared with or separate from Side A/B per policy structure. '
         'Side C for private companies: entity coverage for employment claims, regulatory, '
         'antitrust, and customer claims only.'),
        ('DO-203', 'DIC (Difference in Conditions) Requirements',
         'Side A DIC required for all public companies with market cap > $1B. '
         'DIC conditions: drops down when underlying D&O is exhausted, eroded by defense costs, '
         'insolvent, or fails to respond due to exclusion, rescission, or bankruptcy of insurer. '
         'DIC cannot impose broader exclusions than the underlying D&O policy. '
         'DIC limit: typically 50% to 100% of primary D&O limit.'),
        ('DO-204', 'Hammer Clause Standards',
         'Maximum consent-to-settle hammer ratio: 70/30 (AIG pays 70%, insured bears 30% '
         'of any excess settlement amount beyond the recommended settlement). '
         'Full hammer (100/0): prohibited on AIG Financial Lines policies. '
         'Modified hammer (50/50): available only for large-cap and mega-cap accounts '
         'with established claims history. '
         'All hammer clauses: defense costs continue to be covered regardless of settlement dispute.'),
        ('DO-205', 'Allocation Methodology',
         'For mixed claims (involving both covered and non-covered parties or matters): '
         'Apply the "larger settlement" allocation method as the primary approach. '
         'Relative exposure allocation: acceptable as secondary methodology. '
         'Pre-determined allocation: 80% covered / 20% non-covered for securities class actions '
         'as default unless parties agree otherwise. '
         'Entity vs. individual allocation: 50/50 presumption absent specific facts.'),
        ('DO-206', 'Order of Payments Priority',
         'In the event of policy exhaustion, AIG will apply the following priority of payments: '
         '(1) Side A (individual non-indemnifiable loss) first. '
         '(2) Side B (corporate indemnification) second. '
         '(3) Side C (entity securities/employment) third. '
         'Priority of payments clause is mandatory on all public company D&O placements.'),
    ])

    # ── SECTION 3: SECURITIES CLAIMS ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 3: SECURITIES CLAIMS', level=1)

    add_body(doc,
             'Securities class action claims represent the most significant exposure on public '
             'company D&O policies. These standards govern claim reporting, retroactive date '
             'management, and known-circumstances treatment.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('DO-301', 'Reporting Requirements',
         'Claims must be reported within 15 days from date of service of process or '
         'written demand. Late reporting beyond 30 days: coverage may be denied if AIG '
         'demonstrates prejudice. Circumstances that may give rise to a claim: '
         'may be reported during the policy period with reasonable detail including '
         'the nature of the wrongful act, identity of potential claimants, and '
         'potential quantum of damages.'),
        ('DO-302', 'Known Circumstances Exclusion Standards',
         'Exclude claims arising from facts, circumstances, or wrongful acts known to any '
         'insured person prior to the policy inception date. Standard: "knew or should have '
         'reasonably foreseen" test applies. Application warranty statement required '
         'from CEO, CFO, General Counsel, and Corporate Secretary. '
         'Material misrepresentation in application: grounds for rescission of entire policy.'),
        ('DO-303', 'Retroactive Date Management',
         'New business: Retroactive date equals policy inception date (no prior acts). '
         'Exception: mature companies (10+ years public) with clean loss history may '
         'receive retroactive date matching first D&O policy inception, with prior acts coverage. '
         'Renewal: Maintain existing retroactive date. Never advance the retroactive date '
         'at renewal without VP approval and minimum 20% premium credit to insured.'),
        ('DO-304', 'SEC Investigation Response',
         'Informal SEC inquiry: reportable as circumstance, defense costs coverage triggered '
         'upon receipt of formal order of investigation (FOI). '
         'Formal SEC investigation: defense costs coverage begins upon receipt of Wells Notice '
         'or subpoena, whichever is earlier. Pre-claim inquiry costs sublimit: $500,000 '
         'within the policy limit. SOX Section 304 clawback defense: covered.'),
        ('DO-305', 'Securities Class Action Benchmarks',
         'Average defense cost for securities class action (through trial): $8M-$15M. '
         'Average settlement for securities class action: $20M-$40M (varies by market cap). '
         'Underwriters must ensure adequate total tower limits relative to these benchmarks. '
         'Minimum recommended total D&O tower: 5% of market cap or $25M, whichever is greater.'),
        ('DO-306', 'Derivative Litigation Standards',
         'Shareholder derivative claims: covered under standard D&O policy. '
         'Corporate therapeutic benefit / fee-shifting: covered up to $2M sublimit. '
         'Special Litigation Committee (SLC) costs: covered as defense costs. '
         'Books and records inspection demands (DGCL Section 220): covered under '
         'pre-claim inquiry sublimit.'),
    ])

    # ── SECTION 4: EPLI TREATMENT ──
    add_heading_styled(doc, 'SECTION 4: EMPLOYMENT PRACTICES LIABILITY TREATMENT', level=1)

    add_body(doc,
             'Employment Practices Liability Insurance (EPLI) covers claims by employees '
             'against the company and its directors/officers for wrongful employment practices. '
             'EPLI may be written as a standalone policy or as part of the D&O program.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('DO-401', 'Required EPLI Limits by Employee Count',
         'Employees 1-50: Minimum $1,000,000 per claim / $1,000,000 aggregate. '
         'Employees 51-250: Minimum $2,000,000 per claim / $2,000,000 aggregate. '
         'Employees 251-1,000: Minimum $3,000,000 per claim / $5,000,000 aggregate. '
         'Employees 1,001-5,000: Minimum $5,000,000 per claim / $5,000,000 aggregate. '
         'Employees > 5,000: Minimum $10,000,000 per claim / $10,000,000 aggregate.'),
        ('DO-402', 'Wage and Hour Exclusion Standards',
         'Standard EPLI policy: Wage and hour claims EXCLUDED (including FLSA, state wage laws, '
         'meal/rest period violations, and independent contractor misclassification). '
         'Wage and hour defense costs sublimit: available by endorsement, maximum $500,000 '
         'or 25% of EPLI limit (whichever is less). Defense costs only; no indemnity coverage. '
         'California and New York accounts: require explicit wage/hour risk assessment '
         'due to elevated PAGA and state law exposure.'),
        ('DO-403', 'Third-Party EPLI Requirements',
         'Third-party EPLI (claims by customers, vendors, patients against employees): '
         'Required for all customer-facing businesses with > 100 employees. '
         'Required industries: hospitality, retail, healthcare, financial services, education. '
         'Third-party EPLI limit: Shared with first-party EPLI aggregate. '
         'Third-party EPLI retention: 2x first-party retention.'),
        ('DO-404', 'EPLI Retention Requirements',
         'All EPLI claims: Minimum retention $10,000 per claim for companies < 100 employees. '
         '$25,000 per claim for 100-500 employees. $50,000 per claim for 500-1,000 employees. '
         '$100,000 per claim for > 1,000 employees. '
         'Class/collective action retention: 5x individual claim retention, '
         'minimum $250,000.'),
        ('DO-405', 'Pre-Employment Practice Requirements',
         'All accounts with EPLI: Must have written employee handbook reviewed by '
         'employment counsel within last 24 months. Must have documented anti-harassment '
         'and anti-discrimination training (annual for managers, biennial for all employees). '
         'Accounts without training program: +15% EPLI surcharge. '
         'Accounts with > 3 EEOC charges in 24 months: refer to VP with loss analysis.'),
    ])

    # ── SECTION 5: FIDUCIARY STANDARDS ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 5: FIDUCIARY LIABILITY STANDARDS', level=1)

    add_body(doc,
             'Fiduciary Liability coverage protects plan fiduciaries against claims alleging '
             'breach of fiduciary duty under ERISA. These guidelines address plan types, '
             'emerging risks, and ESOP-specific considerations.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('DO-501', 'Plan Types Requiring Coverage',
         'All ERISA-governed plans require fiduciary coverage including: '
         'Defined benefit pension plans, 401(k) and 403(b) defined contribution plans, '
         'employee stock ownership plans (ESOPs), health and welfare plans, '
         'retiree medical plans, and cafeteria plans (Section 125). '
         'Non-ERISA plans (government, church): fiduciary coverage optional but recommended. '
         'Multiemployer (Taft-Hartley) plans: require separate policy; do not include on '
         'corporate fiduciary policy.'),
        ('DO-502', 'Excessive Fee Litigation Risk Factors',
         'Elevated risk if any of the following apply: Plan assets > $1 billion, '
         'use of proprietary funds > 30% of plan assets, revenue sharing arrangements '
         'with recordkeeper, total plan expense ratio > 75 bps for equity funds, '
         'failure to document investment committee process quarterly, '
         'no RFP for recordkeeping services within last 5 years. '
         'If 3+ risk factors present: apply +25% fiduciary surcharge and '
         'require plan governance questionnaire.'),
        ('DO-503', 'ESOP Valuation Review Requirements',
         'All ESOP risks: require independent annual valuation by qualified ESOP appraiser '
         '(ASA or similar designation). Decline if: ESOP valuation is stale (> 18 months old), '
         'valuation was performed by a non-independent appraiser, or the company has had '
         'a material financial event since last valuation (M&A, restatement, major litigation). '
         'ESOP transaction coverage (leveraged ESOP formation): require separate transaction '
         'liability policy; do not cover under standard fiduciary.'),
        ('DO-504', 'Settlor Function Exclusion',
         'Standard fiduciary policy excludes settlor function activities including: '
         'plan design, amendment, termination, and employer contribution decisions. '
         'Ensure insured understands the distinction between fiduciary and settlor functions. '
         'Settlor function buyback: not available on AIG Fiduciary form.'),
        ('DO-505', 'Voluntary Compliance Programs',
         'Credit for voluntary compliance: If insured participates in DOL Voluntary Fiduciary '
         'Correction Program (VFCP) or IRS Employee Plans Compliance Resolution System (EPCRS): '
         'apply -5% fiduciary premium credit. Require documentation of participation. '
         'Self-correction (SCP) under EPCRS: acceptable for -3% credit.'),
    ])

    # ── SECTION 6: FINANCIAL ANALYSIS ──
    add_heading_styled(doc, 'SECTION 6: FINANCIAL ANALYSIS', level=1)

    add_body(doc,
             'D&O underwriting requires thorough financial analysis of the insured entity. '
             'The following financial health indicators and thresholds guide risk selection '
             'and pricing decisions.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('DO-601', 'D&O Financial Health Indicators',
         'Evaluate the following for all D&O submissions: '
         'Altman Z-Score: Decline if < 1.23 (distress zone). Refer if 1.23-2.90 (grey zone). '
         'Accept if > 2.90 (safe zone). Current Ratio: Minimum 1.0 required; < 0.8 = decline. '
         'Debt-to-Equity Ratio: Maximum 3.0; > 4.0 = decline. '
         'Operating Cash Flow: Must be positive in 2 of last 3 fiscal years. '
         'Going Concern Opinion: Auto-decline if auditor has issued a going concern qualification.'),
        ('DO-602', 'Bankruptcy Risk Thresholds',
         'Companies in Chapter 11 reorganization: decline all new D&O placements. '
         'Companies emerged from bankruptcy < 3 years: Tier 4 risk classification, '
         'maximum $5M D&O limit, retroactive date = emergence date. '
         'Companies with publicly traded debt trading below 60 cents on the dollar: '
         'refer to VP with bankruptcy probability analysis. '
         'Pre-petition D&O tail: available for 6-year extended reporting period at 200% of premium.'),
        ('DO-603', 'M&A Disclosure Requirements',
         'All pending or completed M&A transactions within last 24 months must be disclosed. '
         'Completed acquisitions: require target company loss history (5 years). '
         'Target company integration: covered under acquiring company\'s policy after 90-day '
         'reporting window; longer integration periods require endorsement. '
         'Pending transactions: full transaction details required including advisors, '
         'consideration structure, regulatory approvals, and timeline. '
         'Hostile takeovers: do not provide target-side run-off coverage.'),
        ('DO-604', 'Restatement Analysis',
         'Any financial restatement within last 5 years: auto-refer to VP. '
         'Revenue recognition restatement: +30% D&O surcharge minimum. '
         'Material weakness in internal controls: +20% surcharge and require remediation plan. '
         'SEC comment letter regarding accounting: document resolution before quoting. '
         'Multiple restatements in 5 years: decline.'),
        ('DO-605', 'Stock Volatility Analysis',
         'For public companies: calculate 1-year stock price volatility (beta). '
         'Beta > 2.0: elevated D&O risk, apply +15% surcharge. '
         'Stock price decline > 30% in any 90-day period during last 12 months: '
         'require securities litigation risk assessment before quoting. '
         'Penny stock (share price < $1.00): decline or refer to VP with justification.'),
    ])

    # ── SECTION 7: CLAIM HISTORY ──
    doc.add_page_break()
    add_heading_styled(doc, 'SECTION 7: CLAIM HISTORY AND PRICING', level=1)

    add_body(doc,
             'D&O claim history analysis is critical to pricing adequacy and risk selection. '
             'These standards address loss ratio thresholds, defense cost benchmarks, and '
             'panel counsel guidelines.')

    add_rule_table(doc, ['Rule ID', 'Requirement', 'Threshold / Standard'], [
        ('DO-701', 'Maximum 5-Year D&O Loss Ratio',
         'Public company D&O: Decline if 5-year loss ratio > 80% (incurred including defense costs). '
         'Private company D&O: Decline if 5-year loss ratio > 70%. '
         'EPLI: Decline if 5-year loss ratio > 65%. '
         'Fiduciary: Decline if 5-year loss ratio > 60%. '
         'Loss ratio calculation must be ground-up, including defense costs, '
         'and net of any subrogation or recovery.'),
        ('DO-702', 'Defense Cost Benchmarks',
         'Securities class action (motion to dismiss stage): $1.5M-$3M expected defense costs. '
         'Securities class action (through discovery): $5M-$8M expected defense costs. '
         'Securities class action (through trial): $8M-$15M expected defense costs. '
         'EPLI single plaintiff: $50K-$150K expected defense costs. '
         'EPLI class/collective action: $500K-$2M expected defense costs. '
         'Fiduciary excessive fee litigation: $2M-$5M expected defense costs. '
         'These benchmarks inform reserve adequacy and limit adequacy analysis.'),
        ('DO-703', 'Panel Counsel Rate Guidelines',
         'AIG approved panel counsel rates (maximum hourly without VP approval): '
         'Senior Partner: $950/hour. Partner: $750/hour. Senior Associate: $550/hour. '
         'Associate: $400/hour. Paralegal: $200/hour. '
         'Non-panel counsel: require prior written approval from AIG Claims; rates capped '
         'at panel counsel rates unless exceptional circumstances documented. '
         'Staffing guidelines: maximum 2 partners and 3 associates per active matter.'),
        ('DO-704', 'Claim Frequency Thresholds',
         'Public company D&O: Decline if > 2 securities claims in any 5-year period. '
         'Private company D&O: Decline if > 3 management liability claims in 5 years. '
         'EPLI: Decline if > 5 employment claims per 1,000 employees per year. '
         'EEOC charges: Refer if > 3 charges in any 24-month period. '
         'Prior claim frequency evaluation must cover all prior D&O carriers, not just AIG.'),
        ('DO-705', 'Renewal Pricing Discipline',
         'Maximum renewal rate decrease: -10% without VP approval, -15% with VP approval. '
         'Rate decreases > 15%: require Financial Lines Leader approval and documented rationale. '
         'No rate decreases if any claim reported in expiring policy period. '
         'Minimum rate increase triggers: any reported claim (+10% minimum), '
         'deteriorating financial health indicators (+15% minimum), '
         'industry sector downgrade (+10% minimum), adverse stock volatility event (+20% minimum).'),
        ('DO-706', 'Extended Reporting Period (Tail) Pricing',
         '1-year tail: 100% of annual premium. '
         '2-year tail: 150% of annual premium. '
         '3-year tail: 175% of annual premium. '
         '6-year tail: 200% of annual premium. '
         'Tail must be elected within 60 days of policy expiration/non-renewal. '
         'Change of control: automatic 6-year run-off at 200% premium '
         '(may be pre-negotiated at inception).'),
    ])

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        'AIG Directors & Officers and Fiduciary Liability Underwriting Guidelines v6.0 | '
        'Effective 01/01/2026 | Supersedes all prior versions. These guidelines are '
        'proprietary and confidential. Distribution outside AIG is strictly prohibited.')
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    run.italic = True

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# MAIN
# ============================================================
if __name__ == '__main__':
    print("Generating underwriting guideline DOCX documents...")
    generate_property_guidelines()
    generate_workers_comp_guidelines()
    generate_excess_umbrella_guidelines()
    generate_do_fiduciary_guidelines()
    print("Done — all guideline documents generated.")
