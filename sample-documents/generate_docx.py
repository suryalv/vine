#!/usr/bin/env python3
"""Generate complex underwriting DOCX documents."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

AIG_BLUE = RGBColor(0x00, 0x20, 0x5B)
DARK_TEXT = RGBColor(0x0F, 0x17, 0x2A)
BODY_TEXT = RGBColor(0x33, 0x41, 0x55)
MUTED = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def style_header_row(row, bg_color='00205B'):
    """Style a table header row."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(9)


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = AIG_BLUE
    return h


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = BODY_TEXT
    run.bold = bold
    p.paragraph_format.space_after = Pt(6)
    return p


# ============================================================
# 1. STATEMENT OF VALUES (SOV)
# ============================================================
def generate_sov():
    path = os.path.join(OUT_DIR, 'Statement_of_Values_All_Locations.docx')
    doc = Document()

    # Title
    title = doc.add_heading('Statement of Values', level=0)
    for run in title.runs:
        run.font.color.rgb = AIG_BLUE
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('AIG Commercial Property  |  Prepared: February 2026  |  Effective: 03/01/2026')
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    doc.add_paragraph()

    # Insured Info
    add_heading_styled(doc, 'Insured Information', level=2)
    info_table = doc.add_table(rows=6, cols=2, style='Table Grid')
    info_table.autofit = True
    info_data = [
        ('Named Insured', 'Meridian Steel Corporation'),
        ('Policy Number', 'AIG-CPP-2026-04871'),
        ('Producer', 'Marsh & McLennan — David Chen, SVP'),
        ('Valuation Basis', 'Replacement Cost (unless noted)'),
        ('Currency', 'USD'),
        ('Total Insured Value', '$148,250,000'),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.cell(i, 0).text = k
        info_table.cell(i, 1).text = v
        for p in info_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in info_table.cell(i, 1).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # Location 1
    add_heading_styled(doc, 'Location 1 — Houston Manufacturing Facility', level=2)
    add_body(doc, '4200 Industrial Parkway, Suite 100, Houston, TX 77041')

    loc1_table = doc.add_table(rows=13, cols=2, style='Table Grid')
    loc1_data = [
        ('Occupancy', 'Steel Manufacturing & Warehouse'),
        ('Construction Type', 'Fire Resistive (ISO Class 6)'),
        ('Year Built', '2008  |  Last Renovated: 2022'),
        ('Total Area', '120,000 sq ft (Main: 95,000 / Office: 25,000)'),
        ('Stories', '1 (Manufacturing) / 2 (Office wing)'),
        ('Roof Type', 'Built-up with modified bitumen membrane — inspected 06/2025'),
        ('Sprinkler System', 'Full ESFR — NFPA 13 compliant, last tested 01/2026'),
        ('Fire Alarm', 'Addressable system — central station monitored (ADT)'),
        ('Electrical', 'Primary: 4,160V / 480V transformer. Emergency generator: 500kW Caterpillar'),
        ('HVAC', '12 rooftop units — replaced 2022. Industrial ventilation + dust collection.'),
        ('Flood Zone', 'Zone X (Minimal Risk) — Community Panel: 48201C0405J'),
        ('Wind Tier', 'Tier 2 — Harris County inland (>10 mi from coast)'),
        ('Distance to Fire Station', '2.4 miles — Houston FD Station #82 (ISO PPC: 2)'),
    ]
    for i, (k, v) in enumerate(loc1_data):
        loc1_table.cell(i, 0).text = k
        loc1_table.cell(i, 1).text = v
        for p in loc1_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in loc1_table.cell(i, 1).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    # Location 1 Values
    add_heading_styled(doc, 'Location 1 — Values Detail', level=3)
    val1 = doc.add_table(rows=9, cols=4, style='Table Grid')
    val1_headers = ['Category', 'Replacement Cost', 'ACV', 'Notes']
    for j, h in enumerate(val1_headers):
        val1.cell(0, j).text = h
    style_header_row(val1.rows[0])

    val1_data = [
        ('Building — Main Structure', '$32,000,000', '$24,800,000', 'Includes foundation, walls, roof'),
        ('Building — Office Wing', '$6,500,000', '$4,200,000', '2-story steel frame'),
        ('Business Personal Property', '$12,400,000', '$8,600,000', 'Machinery, raw materials, finished goods'),
        ('CNC Machinery & Equipment', '$8,200,000', '$5,100,000', '6 CNC mills, 4 lathes, 2 plasma cutters'),
        ('Electrical & HVAC Systems', '$3,800,000', '$2,400,000', 'Transformer, generator, HVAC units'),
        ('EDP Equipment', '$1,200,000', '$800,000', 'Servers, network, CAD workstations'),
        ('Improvements & Betterments', '$1,800,000', '$1,200,000', '2022 renovation — upgraded loading docks'),
        ('Outdoor Property & Signage', '$400,000', '$250,000', 'Parking lot, fencing, signage'),
    ]
    for i, row_data in enumerate(val1_data):
        for j, val in enumerate(row_data):
            val1.cell(i+1, j).text = val
            for p in val1.cell(i+1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)

    add_body(doc, 'Location 1 Total Insured Value: $66,300,000', bold=True)

    doc.add_paragraph()

    # Location 2
    add_heading_styled(doc, 'Location 2 — Beaumont Distribution Center', level=2)
    add_body(doc, '8901 Port Arthur Road, Beaumont, TX 77705')

    loc2_table = doc.add_table(rows=13, cols=2, style='Table Grid')
    loc2_data = [
        ('Occupancy', 'Distribution & Light Assembly'),
        ('Construction Type', 'Masonry Non-Combustible (ISO Class 4)'),
        ('Year Built', '2015'),
        ('Total Area', '65,000 sq ft (Warehouse: 55,000 / Office: 10,000)'),
        ('Stories', '1'),
        ('Roof Type', 'Standing seam metal — rated for 130 mph wind load'),
        ('Sprinkler System', 'Full Wet Pipe — NFPA 13 compliant'),
        ('Fire Alarm', 'Conventional zones — central station monitored'),
        ('Electrical', 'Primary: 480V service. UPS for IT systems.'),
        ('Flood Zone', 'Zone AE (High Risk) — BFE: 14.0 ft, Lowest Floor: 16.4 ft'),
        ('Wind Tier', 'Tier 1 — Jefferson County coastal'),
        ('Hurricane Shutters', 'Accordion-style on all openings — rated Cat 3'),
        ('Distance to Fire Station', '3.8 miles — Beaumont FD Station #6 (ISO PPC: 3)'),
    ]
    for i, (k, v) in enumerate(loc2_data):
        loc2_table.cell(i, 0).text = k
        loc2_table.cell(i, 1).text = v
        for p in loc2_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in loc2_table.cell(i, 1).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()

    add_heading_styled(doc, 'Location 2 — Values Detail', level=3)
    val2 = doc.add_table(rows=7, cols=4, style='Table Grid')
    for j, h in enumerate(val1_headers):
        val2.cell(0, j).text = h
    style_header_row(val2.rows[0])

    val2_data = [
        ('Building', '$14,200,000', '$11,600,000', 'Masonry construction'),
        ('Business Personal Property', '$6,800,000', '$4,900,000', 'Inventory, packaging equipment'),
        ('Conveyor & Sorting Systems', '$3,400,000', '$2,100,000', 'Automated conveyor installed 2020'),
        ('Loading Dock Equipment', '$1,200,000', '$850,000', '8 dock levelers, truck restraints'),
        ('EDP & Communication', '$450,000', '$300,000', 'WMS system, RF scanners'),
        ('Outdoor Property', '$350,000', '$220,000', 'Truck yard, fencing, lighting'),
    ]
    for i, row_data in enumerate(val2_data):
        for j, val in enumerate(row_data):
            val2.cell(i+1, j).text = val
            for p in val2.cell(i+1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)

    add_body(doc, 'Location 2 Total Insured Value: $26,400,000', bold=True)

    doc.add_paragraph()

    # Business Income
    add_heading_styled(doc, 'Business Income Worksheet', level=2)
    bi_table = doc.add_table(rows=10, cols=3, style='Table Grid')
    bi_headers = ['Item', 'Location 1', 'Location 2']
    for j, h in enumerate(bi_headers):
        bi_table.cell(0, j).text = h
    style_header_row(bi_table.rows[0])

    bi_data = [
        ('Annual Gross Revenue', '$68,400,000', '$24,200,000'),
        ('Cost of Goods Sold (non-continuing)', '$38,700,000', '$15,100,000'),
        ('Gross Earnings', '$29,700,000', '$9,100,000'),
        ('Payroll (continuing)', '$12,400,000', '$4,800,000'),
        ('Estimated Period of Restoration', '6-9 months', '4-6 months'),
        ('Maximum Foreseeable Loss (BI)', '$22,000,000', '$6,800,000'),
        ('Selected BI Limit', '$18,000,000', '$6,000,000'),
        ('Extra Expense Estimate', '$2,000,000', '$800,000'),
        ('Combined BI + EE Limit', '$20,000,000', '$6,800,000'),
    ]
    for i, row_data in enumerate(bi_data):
        for j, val in enumerate(row_data):
            bi_table.cell(i+1, j).text = val
            for p in bi_table.cell(i+1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)
                    if j == 0:
                        run.bold = True

    doc.add_paragraph()
    add_body(doc, 'Grand Total — All Locations TIV: $148,250,000', bold=True)
    add_body(doc, '(Buildings: $52,700,000 + BPP: $34,800,000 + Machinery/Equipment: $16,850,000 + BI/EE: $26,800,000 + Other: $17,100,000)')

    p = doc.add_paragraph()
    run = p.add_run('\nThis Statement of Values is provided in conjunction with the insurance application and is incorporated by reference into the policy. The insured certifies that the values reported herein are accurate to the best of their knowledge as of the preparation date.')
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# 2. ENDORSEMENT PACKAGE
# ============================================================
def generate_endorsement_package():
    path = os.path.join(OUT_DIR, 'Endorsement_Package_Wind_Hail_Amended.docx')
    doc = Document()

    title = doc.add_heading('Endorsement Package', level=0)
    for run in title.runs:
        run.font.color.rgb = AIG_BLUE

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Wind & Hail Coverage — Amended Terms\nPolicy: AIG-CPP-2026-04871  |  Insured: Meridian Steel Corporation\nEffective: 03/01/2026')
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    doc.add_paragraph()

    # Endorsement 1
    add_heading_styled(doc, 'Endorsement AIG-WH-001: Wind/Hail Percentage Deductible', level=2)
    add_body(doc, 'This endorsement modifies the Commercial Property Coverage Form to which it is attached.')
    add_body(doc, 'SCHEDULE:')

    ded_table = doc.add_table(rows=3, cols=4, style='Table Grid')
    ded_headers = ['Location', 'Wind Tier', 'Deductible %', 'Minimum Deductible']
    for j, h in enumerate(ded_headers):
        ded_table.cell(0, j).text = h
    style_header_row(ded_table.rows[0])
    ded_data = [
        ('Loc 1 — Houston', 'Tier 2', '2% of TIV at time of loss', '$250,000'),
        ('Loc 2 — Beaumont', 'Tier 1', '5% of TIV at time of loss', '$500,000'),
    ]
    for i, row_data in enumerate(ded_data):
        for j, val in enumerate(row_data):
            ded_table.cell(i+1, j).text = val
            for p in ded_table.cell(i+1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    provisions = [
        ('A. Applicable Perils', 'This percentage deductible applies to direct physical loss or damage caused by or resulting from wind, hurricane, tornado, hail, or windborne debris, including but not limited to: (1) the peril of windstorm or hail as defined in the policy; (2) storm surge, tidal wave, or flood directly caused by a named storm; and (3) loss or damage caused by rain, snow, sleet, or dust driven by wind through openings created by the covered peril.'),
        ('B. Deductible Calculation', 'The deductible shall be calculated as the stated percentage of the Total Insured Value (TIV) at the affected location at the time of loss, as reported in the most recent Statement of Values on file with AIG. The deductible applies separately to each location and separately to each occurrence. In no event shall the deductible be less than the minimum deductible shown in the schedule above.'),
        ('C. Named Storm Definition', 'For purposes of this endorsement, a "Named Storm" is any storm or weather disturbance that has been declared and named by the National Hurricane Center (NHC) or the National Weather Service (NWS) as a tropical depression, tropical storm, or hurricane. Named storm deductible shall apply beginning 48 hours prior to the issuance of a tropical storm warning for the area where the insured location is situated and ending 72 hours after the warning is lifted.'),
        ('D. Application to Business Income', 'The percentage deductible also applies to any Business Income and Extra Expense loss resulting from the covered wind/hail peril. The Business Income waiting period specified in the Declarations page applies in addition to (not in lieu of) this percentage deductible.'),
        ('E. Anti-Concurrent Causation', 'When a wind/hail loss occurs concurrently with a flood event (whether or not the flood is a covered peril), the loss shall be adjusted under this endorsement for the wind/hail component and under the flood provisions for the flood component. If the respective components cannot be separated, the entire loss shall be subject to the higher of the wind/hail or flood deductible.'),
    ]

    for heading, body in provisions:
        add_heading_styled(doc, heading, level=3)
        add_body(doc, body)

    doc.add_page_break()

    # Endorsement 2
    add_heading_styled(doc, 'Endorsement AIG-WH-002: Windstorm Protective Safeguards', level=2)
    add_body(doc, 'This endorsement modifies the Commercial Property Coverage Form to which it is attached.')
    add_body(doc, 'As a condition of coverage, the Named Insured warrants compliance with the following windstorm protective safeguard requirements:', bold=True)

    safeguards = [
        '1. HURRICANE PREPAREDNESS PLAN: The insured shall maintain a written Hurricane Preparedness Plan that is reviewed and updated annually no later than May 1st of each year. The plan shall include: (a) designation of a storm preparation team, (b) inventory of emergency supplies and materials, (c) procedures for securing or relocating high-value equipment and inventory, (d) communication protocol for employees and critical vendors, and (e) post-storm damage assessment procedures.',
        '2. ROOF MAINTENANCE: All roof systems shall be inspected by a licensed roofing contractor at least semi-annually (recommended: March and September). Inspection reports shall document: membrane condition, flashing integrity, fastener patterns, ponding water, and drainage capacity. Reports shall be retained for a minimum of 5 years and provided to AIG upon request.',
        '3. OPENING PROTECTION (LOCATION 2): All exterior openings at Location 2 (Beaumont, TX — Wind Tier 1) shall be protected by impact-resistant shutters or panels rated for minimum design wind speed of 130 mph per ASCE 7-22. Shutters shall be deployed and secured when the National Weather Service issues a Hurricane Watch for Jefferson County.',
        '4. BACKUP POWER: Emergency generators at all locations shall be tested monthly under load and serviced quarterly per manufacturer specifications. Fuel reserves shall be maintained at a minimum of 72 hours of continuous operation.',
        '5. TREE & DEBRIS MANAGEMENT: Trees within 50 feet of any insured structure shall be trimmed annually to remove dead branches and reduce windborne debris exposure. All exterior materials and equipment shall be properly secured or stored inside when a Tropical Storm or Hurricane Warning is issued.',
    ]
    for s in safeguards:
        add_body(doc, s)

    doc.add_paragraph()
    add_body(doc, 'IMPORTANT: Failure to comply with the above safeguards may result in denial of a wind/hail claim or a reduction in loss payment as determined by AIG at the time of loss adjustment.', bold=True)

    doc.add_page_break()

    # Endorsement 3
    add_heading_styled(doc, 'Endorsement AIG-WH-003: Storm Surge Buyback', level=2)
    add_body(doc, 'This endorsement amends the Water Exclusion in the underlying policy form.')
    doc.add_paragraph()

    add_heading_styled(doc, 'A. Coverage Provided', level=3)
    add_body(doc, 'Notwithstanding the Water Exclusion (Form CP 10 30 or equivalent) in the underlying policy, this endorsement provides limited coverage for storm surge as follows:')
    add_body(doc, 'Storm surge is defined as the abnormal rise of water generated by a Named Storm (as defined in Endorsement AIG-WH-001), over and above the predicted astronomical tide.')

    surge_table = doc.add_table(rows=3, cols=3, style='Table Grid')
    surge_headers = ['Location', 'Storm Surge Limit', 'Storm Surge Deductible']
    for j, h in enumerate(surge_headers):
        surge_table.cell(0, j).text = h
    style_header_row(surge_table.rows[0])
    surge_data = [
        ('Loc 1 — Houston', '$2,000,000', '$250,000'),
        ('Loc 2 — Beaumont', '$1,000,000', '$500,000'),
    ]
    for i, row_data in enumerate(surge_data):
        for j, val in enumerate(row_data):
            surge_table.cell(i+1, j).text = val
            for p in surge_table.cell(i+1, j).paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()

    add_heading_styled(doc, 'B. Exclusions', level=3)
    exclusions = [
        'This storm surge buyback does NOT cover:',
        '• Flooding from rivers, streams, lakes, or other bodies of fresh water, regardless of whether caused by a Named Storm.',
        '• Surface water runoff or accumulation of water from rainfall.',
        '• Mudflow, mudslide, or earth movement regardless of cause.',
        '• Water that backs up through sewers or drains.',
        '• Water below the surface of the ground, including water which exerts pressure on or seeps through foundations, walls, or floors.',
    ]
    for e in exclusions:
        add_body(doc, e)

    add_heading_styled(doc, 'C. Additional Premium', level=3)
    add_body(doc, 'Annual Premium for Storm Surge Buyback: $12,500')
    add_body(doc, 'This premium is included in the total policy premium shown on the Declarations page.')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ALL OTHER TERMS AND CONDITIONS REMAIN UNCHANGED.')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = AIG_BLUE

    doc.save(path)
    print(f"  Created: {path}")


# ============================================================
# 3. UNDERWRITING SUBMISSION SUMMARY
# ============================================================
def generate_uw_submission_summary():
    path = os.path.join(OUT_DIR, 'UW_Submission_Summary_Greenfield_Agri.docx')
    doc = Document()

    title = doc.add_heading('Underwriting Submission Summary', level=0)
    for run in title.runs:
        run.font.color.rgb = AIG_BLUE

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('CONFIDENTIAL — For Internal Underwriting Use Only\nPrepared by: UW Companion AI  |  Date: February 15, 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED

    doc.add_paragraph()

    # Executive Summary
    add_heading_styled(doc, 'Executive Summary', level=2)
    add_body(doc, 'This submission summary has been prepared by the UW Companion AI system for the environmental liability new business submission from Greenfield Agricultural Holdings, Inc. The submission was received from Lockton Companies on February 7, 2026, and requests coverage effective April 1, 2026.')

    add_body(doc, 'AI RISK ASSESSMENT: HIGH RISK — Manual review required before quotation.', bold=True)

    doc.add_paragraph()

    # Key Metrics
    add_heading_styled(doc, 'Submission Overview', level=2)
    overview_table = doc.add_table(rows=12, cols=2, style='Table Grid')
    overview_data = [
        ('Submission ID', 'SUB-1039'),
        ('Named Insured', 'Greenfield Agricultural Holdings, Inc.'),
        ('DBA', 'Greenfield Farms, Greenfield Organics, Valley Fresh Produce'),
        ('Line of Business', 'Environmental Liability (Site Pollution + Contractors Pollution)'),
        ('Requested Limit', '$5,000,000 per incident / $10,000,000 aggregate'),
        ('Requested Deductible', '$50,000 per incident'),
        ('Proposed Premium', '$560,000 (estimated)'),
        ('Policy Period', '04/01/2026 — 04/01/2029 (3-year term requested)'),
        ('Producer', 'Lockton Companies — Austin, TX\nAccount Executive: James Whitfield'),
        ('Prior Carrier', 'Great American Insurance — declining to renew'),
        ('Reason for Marketing', 'Non-renewal by incumbent due to emerging PFAS contamination exposure'),
        ('AI Confidence Score', '72/100 (Moderate — insufficient historical data on PFAS exposure)'),
    ]
    for i, (k, v) in enumerate(overview_data):
        overview_table.cell(i, 0).text = k
        overview_table.cell(i, 1).text = v
        for p in overview_table.cell(i, 0).paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        for p in overview_table.cell(i, 1).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

    doc.add_page_break()

    # Risk Analysis
    add_heading_styled(doc, 'AI-Generated Risk Analysis', level=2)

    risk_items = [
        ('1. PFAS Contamination Exposure (CRITICAL)',
         'Greenfield operates 4 agricultural processing facilities in Central Texas. Phase II Environmental Site Assessments (ESAs) from 2024 indicate detectable levels of PFAS compounds (specifically PFOA and PFOS) in groundwater monitoring wells at 2 of 4 facilities. Concentrations range from 12-28 ppt (parts per trillion), which currently fall below the EPA\'s proposed MCL of 4 ppt for PFOA and PFOS individually when measured in drinking water sources. However, the onsite levels exceed several state advisory levels and represent a significant emerging risk.\n\nThe source of contamination has been preliminarily linked to the use of PFAS-containing biosolids as agricultural soil amendments between 2015-2022. Greenfield ceased using biosolids in 2022 after the Texas Commission on Environmental Quality (TCEQ) issued an advisory. However, PFAS compounds are persistent and remediation costs are uncertain.'),

        ('2. Regulatory Environment (HIGH)',
         'The EPA finalized the PFAS National Primary Drinking Water Regulation in April 2024, establishing enforceable MCLs for 6 PFAS compounds. Texas has not yet adopted state-specific PFAS standards but TCEQ has been actively monitoring agricultural operations. Greenfield is currently subject to a TCEQ Compliance Investigation (Case #CI-2025-14782) related to groundwater monitoring requirements.\n\nAdditionally, the Comprehensive Environmental Response, Compensation, and Liability Act (CERCLA) designation of PFOA and PFOS as hazardous substances (effective July 2026) could expose Greenfield to Superfund liability for any contamination that migrates off-site.'),

        ('3. Third-Party Bodily Injury Risk (MODERATE)',
         'Three neighboring residential properties are served by private water wells within a 1-mile radius of the affected facilities. While current testing shows no PFAS detection in these wells, the potential for groundwater migration exists. Two neighboring property owners have retained counsel (Rodriguez & Associates) and sent preservation of evidence letters to Greenfield. No formal claims have been filed to date.'),

        ('4. Remediation Cost Uncertainty (HIGH)',
         'Environmental consultants (Tetra Tech) have estimated remediation costs ranging from $2.8M to $12.4M depending on the scope and method of cleanup:\n\n• Monitored Natural Attenuation: $2.8M over 10 years (least aggressive)\n• Pump-and-Treat with GAC Filtration: $5.2M over 5-7 years\n• In-Situ Soil Treatment + Groundwater: $8.7M over 3-5 years\n• Full Excavation & Off-Site Disposal: $12.4M over 2-3 years\n\nThe remediation approach has not been finalized pending the outcome of the TCEQ investigation.'),

        ('5. Financial Capacity (MODERATE)',
         'Greenfield Agricultural Holdings reported FY 2025 revenue of $142M with net income of $8.7M. Debt-to-equity ratio is 0.84. The company has $12M in available credit facilities. While the balance sheet is generally healthy, a worst-case remediation scenario ($12.4M) would represent significant financial stress.\n\nAIG Financial Analysis Unit recommends requiring audited financial statements and a letter of credit for the deductible amount ($50,000) as a condition of binding.'),
    ]

    for heading, body in risk_items:
        add_heading_styled(doc, heading, level=3)
        add_body(doc, body)

    doc.add_page_break()

    # AI Recommendations
    add_heading_styled(doc, 'AI Underwriting Recommendations', level=2)

    recs = [
        ('Coverage Modifications',
         '1. Reduce per-incident limit from $5M to $3M given PFAS uncertainty\n2. Apply PFAS-specific sublimit of $1M per incident / $2M aggregate\n3. Increase deductible to $100K per incident (minimum)\n4. Exclude pre-existing conditions at Facilities 2 and 3 (where PFAS detected)\n5. Add regulatory defense cost sublimit of $500K\n6. Policy term: 1-year only (not 3-year as requested) — reassess at renewal'),
        ('Required Documentation Before Quoting',
         '1. Complete Phase II ESA reports for all 4 facilities (2024 reports on file)\n2. Groundwater monitoring data from last 4 quarters\n3. TCEQ Compliance Investigation correspondence and status\n4. Audited financial statements (FY 2024 and FY 2025)\n5. Copy of remediation consultant\'s scope of work proposal\n6. Environmental compliance history for past 10 years\n7. Third-party claim correspondence (Rodriguez & Associates letters)'),
        ('Pricing Guidance',
         'Base Rate: $560,000 (as quoted by producer)\nAI-Recommended Rate: $685,000 - $740,000\n\nRate Adjustments:\n• +15% for PFAS exposure surcharge\n• +8% for active regulatory investigation\n• +5% for 3rd party BI potential\n• -3% credit for proactive biosolid cessation (2022)\n• -2% credit for environmental consultant engagement\n\nMinimum Acceptable Premium: $685,000 for reduced limits as recommended above.\nIf full $5M limits are maintained: Minimum premium $920,000 with PFAS exclusion.'),
    ]

    for heading, body in recs:
        add_heading_styled(doc, heading, level=3)
        add_body(doc, body)

    doc.add_paragraph()

    # Conclusion
    add_heading_styled(doc, 'Conclusion & Next Steps', level=2)
    add_body(doc, 'The UW Companion AI system has flagged this submission as HIGH RISK due to the emerging PFAS contamination exposure, active regulatory investigation, and potential third-party claims. While Greenfield Agricultural Holdings has demonstrated proactive risk management by ceasing biosolid use and engaging environmental consultants, the uncertainty around remediation costs and regulatory outcomes warrants careful underwriting consideration.')
    add_body(doc, 'Recommended Action: Refer to Senior Environmental Underwriter for manual review. Do not issue quote until all required documentation is received and reviewed.', bold=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('This report was generated by UW Companion AI. All analysis and recommendations are advisory and should be reviewed by a qualified underwriter before any binding decisions are made. AI confidence score: 72/100.')
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED
    run.italic = True

    doc.save(path)
    print(f"  Created: {path}")


if __name__ == '__main__':
    print("Generating DOCX documents...")
    generate_sov()
    generate_endorsement_package()
    generate_uw_submission_summary()
    print("Done — all DOCX documents generated.")
